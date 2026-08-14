"""SOT-2703 — 複合書式ファクトストア + 決定論 lookup レーン（cycle10, idx16/71）の offline テスト.

ネットワーク/LLM/実コーパス非依存。in-memory python-docx で run 書式抽出器（highlight/font_color/
bold/underline/italic）を、合成ストア（``format_facts_store.load`` を monkeypatch）で 2 本の決定論
束縛（idx16 黄×赤 の doc-kind 自己同定による一意化 / idx71 太字∧下線∧イタリックの案件横断一意化）と
精度優先の deferral、``RAG_FORMAT_FACTS`` 既定 OFF の byte-identical 挙動を検証する。
"""
from __future__ import annotations

import io

import docx
import pytest
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

from src.rag.agent import fact_layer as fl
from src.rag.agent import format_facts_lane as L
from src.rag.index import format_facts_store as S
from src.rag.corpus import FileRef


# =========================================================================== run 書式抽出（in-memory docx）
def test_run_attrs_extracts_all_five(monkeypatch):
    """黄ハイライト×赤字 run と 太字∧下線∧イタリック run を 1 文書から全属性で拾う。"""
    d = docx.Document()
    p = d.add_paragraph()
    r1 = p.add_run("0.589")
    r1.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r1.font.color.rgb = RGBColor(0xEE, 0x00, 0x00)
    r2 = p.add_run("29290")               # 黄ハイライトのみ（赤字ではない）
    r2.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p2 = d.add_paragraph()
    r3 = p2.add_run("4,250,000円")
    r3.bold = True
    r3.underline = True
    r3.italic = True
    p2.add_run("（税抜）")                 # 書式なし run は捨てられる

    buf = io.BytesIO()
    d.save(buf)
    data = buf.getvalue()
    ref = FileRef(path="/tmp/x.docx", rel="プロジェクト/X/05.会議/報告資料/報告資料_2025-04-29.docx",
                  name="報告資料_2025-04-29.docx", project="X", ext="docx", category="report")
    monkeypatch.setattr(S, "_office_bytes", lambda r: data)

    runs = S._doc_runs(ref)
    texts = {r["text"]: r["attrs"] for r in runs}
    assert set(texts) == {"0.589", "29290", "4,250,000円"}      # 書式なし「（税抜）」は不収蔵
    assert texts["0.589"]["highlight"] == "黄" and texts["0.589"]["font_color"] == "赤"
    assert texts["29290"]["highlight"] == "黄" and texts["29290"]["font_color"] is None
    bui = texts["4,250,000円"]
    assert bui["bold"] and bui["underline"] and bui["italic"]
    assert bui["highlight"] is None and bui["font_color"] is None


def test_folder_and_selfid_kinds():
    assert S._folder_kinds("プロジェクト/X/05.会議/報告資料/報告資料_2025-04-29.docx",
                           "報告資料_2025-04-29.docx") == ["報告資料"]
    assert S._folder_kinds("プロジェクト/X/05.会議/会議録/会議録_2025-08-06.docx",
                           "会議録_2025-08-06.docx") == ["会議録"]
    # 本文に『中間報告』を含む報告資料 ⇒ 中間報告 を doc-kind に付与（会議録には付けない）。
    assert S._selfid_kinds("本報告は M02 時点の中間報告です。", ["報告資料"]) == ["中間報告"]
    assert S._selfid_kinds("本報告は M02 時点の中間報告です。", ["会議録"]) == []
    assert S._selfid_kinds("キックオフ会。", ["報告資料"]) == []


# =========================================================================== 属性パース
def test_required_attrs_parse():
    assert L._required_attrs(L._norm("黄色ハイライトかつ赤字となっている部分")) == \
        {"highlight": "黄", "font_color": "赤"}
    assert L._required_attrs(L._norm("太字、下線、イタリックのすべてに該当する箇所")) == \
        {"bold": True, "underline": True, "italic": True}
    assert L._required_attrs(L._norm("オレンジのマーカーが引かれた語")) == {"highlight": "オレンジ"}
    assert L._required_attrs(L._norm("案件名を教えてください")) == {}


def test_run_satisfies():
    yr = {"highlight": "黄", "font_color": "赤", "bold": False, "underline": False, "italic": False}
    assert L._run_satisfies(yr, {"highlight": "黄", "font_color": "赤"})
    assert not L._run_satisfies({**yr, "font_color": None}, {"highlight": "黄", "font_color": "赤"})
    bui = {"highlight": None, "font_color": None, "bold": True, "underline": True, "italic": True}
    assert L._run_satisfies(bui, {"bold": True, "underline": True, "italic": True})
    assert not L._run_satisfies({**bui, "italic": False}, {"bold": True, "underline": True, "italic": True})


# =========================================================================== 合成ストア + レーン
def _yr(text):
    return {"text": text, "attrs": {"highlight": "黄", "font_color": "赤",
                                    "bold": False, "underline": False, "italic": False}, "loc": "p"}


def _bui(text, *, bold=True, underline=True, italic=True):
    return {"text": text, "attrs": {"highlight": None, "font_color": None,
                                    "bold": bold, "underline": underline, "italic": italic}, "loc": "p"}


def _records():
    aym = "青葉与信マネジメント株式会社"
    aomine = "株式会社青嶺不動産アセットマネジメント"
    return [
        # idx16: 2 本の報告資料が両方 黄×赤 を持つ。中間報告の自己同定で 04-29 に一意化する必要がある。
        {"project": aym, "rel": f"プロジェクト/{aym}/05.会議/報告資料/報告資料_2025-04-29.docx",
         "name": "報告資料_2025-04-29.docx", "doc_kind": ["報告資料", "中間報告"], "n_runs": 2,
         "runs": [_yr("0.589"),
                  {"text": "29290", "attrs": {"highlight": "黄", "font_color": None, "bold": False,
                                              "underline": False, "italic": False}, "loc": "p"}]},
        {"project": aym, "rel": f"プロジェクト/{aym}/05.会議/報告資料/報告資料_2025-04-09.docx",
         "name": "報告資料_2025-04-09.docx", "doc_kind": ["報告資料"], "n_runs": 1,
         "runs": [_yr("4,620,000")]},
        # idx71: 3 本の会議録のうち 1 本だけ 太字∧下線∧イタリック run を持つ。
        {"project": aomine, "rel": f"プロジェクト/{aomine}/05.会議/会議録/会議録_2025-08-06.docx",
         "name": "会議録_2025-08-06.docx", "doc_kind": ["会議録"], "n_runs": 1,
         "runs": [_bui("4,250,000円")]},
        {"project": aomine, "rel": f"プロジェクト/{aomine}/05.会議/会議録/会議録_2025-09-16.docx",
         "name": "会議録_2025-09-16.docx", "doc_kind": ["会議録"], "n_runs": 1,
         "runs": [_bui("3,000,000円", italic=False)]},   # 全三種を満たさない
        {"project": aomine, "rel": f"プロジェクト/{aomine}/05.会議/会議録/会議録_2025-08-26.docx",
         "name": "会議録_2025-08-26.docx", "doc_kind": ["会議録"], "n_runs": 0, "runs": []},
    ]


@pytest.fixture
def store(monkeypatch):
    monkeypatch.setattr(S, "load", lambda path=None: _records())
    monkeypatch.setenv("RAG_FORMAT_FACTS", "1")
    monkeypatch.setenv("RAG_FACT_LAYER", "1")
    return None


def test_idx16_yellow_red_selfid_binding(store):
    r = L.resolve("青葉与信マネジメントの中間報告資料にて、黄色ハイライトかつ赤字となっている部分を抜き出してください。")
    assert r is not None and r["value"] == "0.589"          # 04-09 の 4,620,000 でなく 中間報告 04-29
    assert "報告資料_2025-04-29" in r["evidence"]["doc_id"]
    assert r["method"]["selection"] == "composite_format"


def test_idx16_via_fact_layer(store):
    r = fl.resolve("青葉与信マネジメントの中間報告資料にて、黄色ハイライトかつ赤字となっている部分を抜き出してください。",
                   "simple_lookup")
    assert r is not None and r["value"] == "0.589"


def test_idx71_bold_underline_italic_binding(store):
    r = L.resolve("青嶺不動産アセットマネジメントの会議録の中で、太字、下線、イタリックのすべてに該当する箇所を抽出してください。")
    assert r is not None and r["value"] == "4,250,000円"
    assert "会議録_2025-08-06" in r["evidence"]["doc_id"]


def test_defers_without_doc_kind_when_ambiguous(store):
    # doc-kind を名指さず 黄×赤 だけ ⇒ 案件内 2 doc の 0.589 と 4,620,000 が両立 ⇒ 曖昧 ⇒ defer。
    assert L.resolve("青葉与信マネジメントで黄色ハイライトかつ赤字の部分を抜き出してください。") is None


def test_defers_unknown_project(store):
    assert L.resolve("存在しない会社の会議録で太字下線イタリックの箇所は？") is None


def test_defers_wrong_color(store):
    # オレンジ×赤字は該当 run が無い ⇒ defer（wrong を増やさない）。
    assert L.resolve("青葉与信マネジメントの中間報告資料でオレンジのハイライトかつ赤字の部分は？") is None


# --------------------------------------------------------------------------- OFF byte-identical
def test_off_is_inert(monkeypatch):
    monkeypatch.setattr(S, "load", lambda path=None: _records())
    monkeypatch.delenv("RAG_FORMAT_FACTS", raising=False)
    assert L.enabled() is False
    assert L.resolve("青葉与信マネジメントの中間報告資料にて、黄色ハイライトかつ赤字の部分を抜き出してください。") is None


def test_off_lane_not_in_fact_layer(monkeypatch):
    monkeypatch.setattr(S, "load", lambda path=None: _records())
    monkeypatch.setenv("RAG_FACT_LAYER", "1")
    monkeypatch.delenv("RAG_FORMAT_FACTS", raising=False)
    assert fl.resolve("青嶺不動産アセットマネジメントの会議録の中で、太字、下線、イタリックのすべての箇所を抽出してください。",
                      "simple_lookup") is None


# --------------------------------------------------------------------------- store schema roundtrip
def test_store_load_schema_roundtrip(tmp_path):
    p = tmp_path / "format_facts_store.jsonl"
    S.write_store(_records(), path=p)
    S.reset_cache()
    rows = S.load(path=p)
    assert len(rows) == 5
    aym = [r for r in rows if "青葉与信" in r["project"]]
    assert len(aym) == 2 and any("中間報告" in r["doc_kind"] for r in aym)


# =========================================================================== SOT-2711 列挙 / emph_spans
def _emph(text, *, bold=False, underline=False, italic=False, loc="p"):
    return {"text": text, "attrs": {"bold": bold, "underline": underline, "italic": italic}, "loc": loc}


def test_is_date_structural_rule():
    assert L._is_date("2025-09-02") and L._is_date("2025/10/7") and L._is_date("2025年9月2日")
    assert not L._is_date("time_and_materials")
    assert not L._is_date("25,000円／時間")      # 数字を含むが日付ではない
    assert not L._is_date("30分単位")


def test_enum_regex_separates_enumerate_from_conjunction():
    # idx3 型「…をすべて抽出」＝列挙。idx11/71 型「すべてに該当する箇所を抽出」＝装飾連言（非列挙）。
    assert L._ENUM_RE.search(L._norm("日付以外のものをすべて抽出してください。"))
    assert not L._ENUM_RE.search(L._norm("太字、下線、イタリックのすべてに該当する箇所を抽出してください。"))


def _kaede_bold_record():
    """idx3: かえで契約書の太字 merged span（日付 2 種 + 実データ 4 件）。"""
    kaede = "医療法人社団 恒一会 かえで総合病院"
    return {
        "project": kaede, "rel": f"プロジェクト/{kaede}/01.契約/契約書_pw-kaede20250902.docx",
        "name": "契約書_pw-kaede20250902.docx", "doc_kind": ["契約書"], "ext": "docx",
        "n_runs": 0, "runs": [],
        "emph_spans": [
            _emph("2025-09-02", bold=True), _emph("2025-10-07", bold=True),
            _emph("time_and_materials", bold=True),
            _emph("実績工数に基づき、案件完了後に最終成果物の検収を経て一括精算する。", bold=True),
            _emph("30分単位", bold=True), _emph("2025-09-02", bold=True),  # 重複日付
            _emph("25,000円／時間", bold=True)],
    }


@pytest.fixture
def enum_store(monkeypatch):
    monkeypatch.setattr(S, "load", lambda path=None: [_kaede_bold_record()])
    monkeypatch.setenv("RAG_FORMAT_FACTS", "1")
    monkeypatch.setenv("RAG_FACT_LAYER", "1")
    return None


def test_idx3_bold_enumeration_excludes_dates(enum_store):
    r = L.resolve("恒一会 かえで総合病院の契約書において、太字で記載されている箇所のうち、"
                  "日付以外のものをすべて抽出してください。")
    assert r is not None
    assert r["value"] == ("time_and_materials、実績工数に基づき、案件完了後に最終成果物の検収を経て"
                          "一括精算する。、30分単位、25,000円／時間")   # 日付 2 種除外・重複除去・文書順
    assert r["method"]["selection"] == "composite_format_enumerate"


def test_idx3_without_date_exclusion_keeps_dates(enum_store):
    # 「日付以外」が無ければ日付も列挙に含む（構造規則は質問駆動; gold 非依存）。
    r = L.resolve("恒一会 かえで総合病院の契約書の太字箇所をすべて列挙してください。")
    assert r is not None and r["value"].startswith("2025-09-02、2025-10-07、time_and_materials")


def _aomine_pdf_bui_record():
    """idx11: 青嶺 報告資料 PDF の B∧U∧I merged span（pdf は runs 無し・emph_spans のみ）。"""
    aomine = "株式会社青嶺不動産アセットマネジメント"
    return {
        "project": aomine, "rel": f"プロジェクト/{aomine}/05.会議/報告資料/報告資料_2025-08-06.pdf",
        "name": "報告資料_2025-08-06.pdf", "doc_kind": ["報告資料"], "ext": "pdf",
        "n_runs": 0, "runs": [],
        "emph_spans": [_emph("4,675,000円", bold=True, underline=True, italic=True)],
    }


def test_idx11_pdf_bold_underline_italic_single_binding(monkeypatch):
    monkeypatch.setattr(S, "load", lambda path=None: [_aomine_pdf_bui_record()])
    monkeypatch.setenv("RAG_FORMAT_FACTS", "1")
    r = L.resolve("青嶺不動産アセットマネジメントの報告資料の中で、"
                  "太字、下線、イタリックのすべてに該当する箇所を抽出してください。")
    assert r is not None and r["value"] == "4,675,000円"
    assert "報告資料_2025-08-06.pdf" in r["evidence"]["doc_id"]
    assert r["method"]["selection"] == "composite_format"       # 単発（列挙でない）
