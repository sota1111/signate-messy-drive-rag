"""SOT-2585: deterministic OOXML FORMAT_EVENT extraction (Excel cfRule/dxf + Word comments/effective style).

All behaviour lives behind ``RAG_FORMAT_EVENTS`` (default OFF); OFF must be byte-identical to the champion
extract face. ON turns the three ``PARSER_CAPABILITY_MISS`` questions into deterministic evidence:

* idx65 — an Excel conditional-format rule surfaced as a full FORMAT_EVENT (condition + dxf + priority).
* idx16 — a docx run that is BOTH yellow-highlighted AND red-font found by a composite predicate.
* idx49 — a docx comment's *anchor text* recovered via the ``document.xml`` ⇄ ``comments.xml`` XML path.
"""
from __future__ import annotations

import zipfile
from pathlib import Path

import docx
import openpyxl
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from openpyxl.formatting.rule import CellIsRule
from openpyxl.styles import Font, PatternFill
from openpyxl.styles.differential import DifferentialStyle
from openpyxl.formatting.rule import Rule

from src.rag.corpus import FileRef
from src.rag.tools.format_events import comment_anchors, enabled, fmt_notation, format_events


def _ref(path: Path, ext: str) -> FileRef:
    return FileRef(path=path, project="p", category="data", rel=path.name, name=path.name, ext=ext)


# =========================================================================== gating
def test_enabled_flag(monkeypatch) -> None:
    monkeypatch.setenv("RAG_FORMAT_EVENTS", "0")
    assert not enabled()
    monkeypatch.setenv("RAG_FORMAT_EVENTS", "1")
    assert enabled()


# =========================================================================== Excel cfRule / dxf
def _corr_xlsx(path: Path) -> None:
    """A correlation-like sheet with a yellow ``<-0.99`` (with red dxf font) and a red ``>0.99`` rule."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "相関係数"
    ws["B2"] = -0.995      # satisfies < -0.99 → yellow
    ws["B3"] = 0.995       # satisfies > 0.99  → red
    ws["B4"] = 0.10
    yellow = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
    # a yellow fill + red differential font, so dxf.fill AND dxf.font are both exercised
    dxf = DifferentialStyle(fill=yellow, font=Font(color="FF0000"))
    ws.conditional_formatting.add("A1:E5", Rule(type="cellIs", operator="lessThan",
                                                formula=["-0.99"], dxf=dxf, priority=1,
                                                stopIfTrue=True))
    red = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
    ws.conditional_formatting.add("A1:E5", CellIsRule(operator="greaterThan", formula=["0.99"],
                                                      fill=red))
    wb.save(path)


def test_excel_conditional_format_event_full_record(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("RAG_STRUCTURE_STORE", "0")
    path = tmp_path / "corr.xlsx"
    _corr_xlsx(path)
    res = format_events(_ref(path, "xlsx"), kind="conditional_format")
    cf = res["value"]
    yellow_rule = next(it for it in cf if it["method"]["fill_color"] == "黄")
    ev, m = yellow_rule["evidence"], yellow_rule["method"]
    assert yellow_rule["value"] == "セルの値 < -0.99"
    assert m["kind"] == "conditional_format" and m["source"] == "conditional"
    assert m["font_color"] == "赤"                       # dxf.font surfaced
    assert ev["rule_type"] == "cellIs" and ev["operator"] == "lessThan"
    assert ev["formula"] == ["-0.99"]
    assert ev["priority"] == 1 and ev["stop_if_true"] is True
    assert ev["dxf_id"] is not None
    # Level-2 evaluation: the satisfying numeric cell is enumerated (openpyxl cached value)
    assert ev["n_matched"] == 1 and "B2" in ev["matched_cells"]


def test_excel_composite_predicate_direct_cells(tmp_path: Path, monkeypatch) -> None:
    """A directly-styled cell (yellow fill + red font) is found by the fill∧font AND predicate."""
    monkeypatch.setenv("RAG_STRUCTURE_STORE", "0")
    path = tmp_path / "direct.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "対象"
    ws["A1"].fill = PatternFill(start_color="FFFFFF00", end_color="FFFFFF00", fill_type="solid")
    ws["A1"].font = Font(color="FFFF0000")
    ws["A2"] = "無地"
    ws["A3"] = "黄のみ"
    ws["A3"].fill = PatternFill(start_color="FFFFFF00", end_color="FFFFFF00", fill_type="solid")
    wb.save(path)
    hit = format_events(_ref(path, "xlsx"), fill="黄", font_color="赤")["value"]
    assert [it["value"] for it in hit] == ["対象"]        # A3 (yellow only) excluded by the AND


# =========================================================================== Word effective style (idx16)
def test_docx_composite_highlight_and_font(tmp_path: Path, monkeypatch) -> None:
    """idx16: only the run that is BOTH yellow-highlighted AND red-font matches fill=黄, font_color=赤."""
    monkeypatch.setenv("RAG_STRUCTURE_STORE", "0")
    path = tmp_path / "report.docx"
    d = docx.Document()
    p = d.add_paragraph()
    r1 = p.add_run("0.589")                 # yellow highlight + red font  ← the answer
    r1.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r1.font.color.rgb = RGBColor(0xEE, 0x00, 0x00)
    r2 = p.add_run("29290")                 # yellow highlight, black font (not red)
    r2.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r3 = p.add_run("300")                   # red font, no highlight
    r3.font.color.rgb = RGBColor(0xEE, 0x00, 0x00)
    d.save(path)

    hit = format_events(_ref(path, "docx"), fill="黄", font_color="赤")["value"]
    assert [it["value"] for it in hit] == ["0.589"]
    m = hit[0]["method"]
    assert m["fill_color"] == "黄" and m["font_color"] == "赤" and m["source"] == "highlight"
    assert hit[0]["evidence"]["style_origin"] == "direct"


def test_fmt_notation_render(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("RAG_STRUCTURE_STORE", "0")
    path = tmp_path / "n.docx"
    d = docx.Document()
    r = d.add_paragraph().add_run("赤黄")
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    d.save(path)
    it = format_events(_ref(path, "docx"), fill="黄", font_color="赤")["value"][0]
    note = fmt_notation(it)
    assert "font.color=赤" in note and "background=黄" in note
    assert "background.source=highlight" in note and "style.origin=direct" in note


def test_docx_effective_background_highlight_over_shading(tmp_path: Path, monkeypatch) -> None:
    """A run carrying both a highlight and run-shading reports the highlight as effective background."""
    monkeypatch.setenv("RAG_STRUCTURE_STORE", "0")
    path = tmp_path / "shd.docx"
    d = docx.Document()
    run = d.add_paragraph().add_run("重要")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # inject run-level shading (green) directly into rPr
    from docx.oxml.ns import qn
    rpr = run._r.get_or_add_rPr()
    shd = rpr.makeelement(qn("w:shd"), {qn("w:fill"): "00B050"})
    rpr.append(shd)
    d.save(path)
    ev = format_events(_ref(path, "docx"), kind="run")["value"][0]["evidence"]
    assert ev["effective_background"] == "黄" and ev["background_source"] == "highlight"
    assert ev["run_shading"] == "00B050"      # the underlying shading is still recorded separately


# =========================================================================== Word comments (idx49, XML path)
def _docx_with_comment(path: Path, *, anchor: str, body: str, author: str, date: str) -> None:
    """Build a minimal .docx whose ``document.xml`` wraps ``anchor`` in a commentRange linked to a comment.

    python-docx 1.1.x has no comment-authoring API, so the test writes the OOXML parts directly — this is
    exactly the XML-direct correspondence the extractor must recover.
    """
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    CT = "http://schemas.openxmlformats.org/package/2006/content-types"
    document = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="{W}"><w:body><w:p><w:r><w:t>前文</w:t></w:r>
<w:commentRangeStart w:id="7"/><w:r><w:t>{anchor}</w:t></w:r><w:commentRangeEnd w:id="7"/>
<w:r><w:commentReference w:id="7"/></w:r></w:p></w:body></w:document>"""
    comments = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="{W}"><w:comment w:id="7" w:author="{author}" w:date="{date}">
<w:p><w:r><w:t>{body}</w:t></w:r></w:p></w:comment></w:comments>"""
    content_types = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="{CT}">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>
</Types>"""
    root_rels = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="{R}/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    doc_rels = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="{R}/comments" Target="comments.xml"/>
</Relationships>"""
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types)
        z.writestr("_rels/.rels", root_rels)
        z.writestr("word/document.xml", document)
        z.writestr("word/_rels/document.xml.rels", doc_rels)
        z.writestr("word/comments.xml", comments)


def test_docx_comment_anchor_via_xml(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("RAG_STRUCTURE_STORE", "0")
    path = tmp_path / "会議録.docx"
    _docx_with_comment(path, anchor="WBS・進捗管理台帳確定（タスク割振・ガント更新）",
                        body="要確認", author="h.ikeshita", date="2026-03-30T19:02:00Z")
    assert comment_anchors(_ref(path, "docx")) == ["WBS・進捗管理台帳確定（タスク割振・ガント更新）"]
    ev = format_events(_ref(path, "docx"), kind="comment")["value"][0]["evidence"]
    assert ev["anchor_text"] == "WBS・進捗管理台帳確定（タスク割振・ガント更新）"
    assert ev["comment_text"] == "要確認" and ev["author"] == "h.ikeshita"
    assert ev["timestamp"] == "2026-03-30T19:02:00Z"
    assert ev["comment_id"] == "7"


def test_docx_comment_body_fallback_without_range(tmp_path: Path, monkeypatch) -> None:
    """A comment whose range markers are absent still surfaces its body (no anchor)."""
    monkeypatch.setenv("RAG_STRUCTURE_STORE", "0")
    path = tmp_path / "noanchor.docx"
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    CT = "http://schemas.openxmlformats.org/package/2006/content-types"
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml",
                   f'<Types xmlns="{CT}"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                   '<Default Extension="xml" ContentType="application/xml"/>'
                   '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                   '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/></Types>')
        z.writestr("_rels/.rels",
                   f'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="{R}/officeDocument" Target="word/document.xml"/></Relationships>')
        z.writestr("word/document.xml", f'<w:document xmlns:w="{W}"><w:body><w:p><w:r><w:t>本文</w:t></w:r></w:p></w:body></w:document>')
        z.writestr("word/comments.xml", f'<w:comments xmlns:w="{W}"><w:comment w:id="3" w:author="a"><w:p><w:r><w:t>単独コメント</w:t></w:r></w:p></w:comment></w:comments>')
    anchors = comment_anchors(_ref(path, "docx"))
    assert anchors == ["単独コメント"]


# =========================================================================== flag-OFF byte-identical
def test_extract_face_byte_identical_when_off(tmp_path: Path, monkeypatch) -> None:
    """office.extract_docx/xlsx output is identical whether the flag is unset — no __FMT__ leakage."""
    from src.rag.extract import office
    monkeypatch.setenv("RAG_STRUCTURE_STORE", "0")
    monkeypatch.setenv("RAG_FONT_EMPHASIS", "0")
    monkeypatch.setenv("RAG_HIGHLIGHT_EXTRA", "0")

    xlsx = tmp_path / "corr.xlsx"
    _corr_xlsx(xlsx)
    docx_path = tmp_path / "report.docx"
    d = docx.Document()
    run = d.add_paragraph().add_run("0.589")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.font.color.rgb = RGBColor(0xEE, 0x00, 0x00)
    d.save(docx_path)

    monkeypatch.setenv("RAG_FORMAT_EVENTS", "0")
    off_x = office.extract_xlsx(_ref(xlsx, "xlsx"), None)
    off_d = office.extract_docx(_ref(docx_path, "docx"), None)
    assert "__FMT__" not in off_x and "FORMAT_EVENT" not in off_x
    assert "__FMT__" not in off_d and "FORMAT_EVENT" not in off_d

    monkeypatch.setenv("RAG_FORMAT_EVENTS", "1")
    on_x = office.extract_xlsx(_ref(xlsx, "xlsx"), None)
    on_d = office.extract_docx(_ref(docx_path, "docx"), None)
    # OFF is a strict prefix of ON (ON only appends the 【書式イベント】/__FMT__ block)
    assert on_x.startswith(off_x) and "__FMT__" in on_x
    assert on_d.startswith(off_d) and "__FMT__" in on_d


def test_tool_registration_gated(monkeypatch) -> None:
    """format_events is exposed to the investigator only when the flag is on (tool set byte-identical off)."""
    from src.rag.agent import investigator
    from src.rag.tools.profile import CorpusProfile
    profile = CorpusProfile()
    monkeypatch.setenv("RAG_FONT_EMPHASIS", "0")
    monkeypatch.setenv("RAG_FORMAT_EVENTS", "0")
    off = {t.name for t in investigator.build_generic_tools(profile)}
    assert "format_events" not in off
    monkeypatch.setenv("RAG_FORMAT_EVENTS", "1")
    on = {t.name for t in investigator.build_generic_tools(profile)}
    assert "format_events" in on and off < on


def test_routing_prefers_format_events_when_enabled(monkeypatch) -> None:
    from src.rag.agent import routing
    from src.rag.agent import question_contract as qc
    monkeypatch.setenv("RAG_FORMAT_EVENTS", "1")
    c = qc.classify("会議録でコメントがついている部分をそのまま抽出してください")
    assert routing.first_tools_for(c, "会議録でコメントがついている部分をそのまま抽出してください")[0] == "format_events"
    q = "黄色ハイライトかつ赤字となっている部分を抜き出してください"
    assert routing.first_tools_for(qc.classify(q), q)[0] == "format_events"
    monkeypatch.setenv("RAG_FORMAT_EVENTS", "0")
    off_first = routing.first_tools_for(c, "会議録でコメントがついている部分をそのまま抽出してください")
    assert off_first and off_first[0] != "format_events"
