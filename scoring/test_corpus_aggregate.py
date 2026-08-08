"""Offline tests for the全プロジェクト横断 corpus_aggregate tool (no LLM / network) — SOT-2489.

    .venv/bin/python -m pytest scoring/test_corpus_aggregate.py -q

Two layers:

* **Synthetic mini-corpus** (always runs): builds a tiny share-drive-shaped corpus of docx contracts +
  train.csv with known amounts / 着手金 / periods / staff and a hand-built glossary, then pins every
  aggregation op (count / max / min / period-filter), 主略称 normalization, round-up, and the
  precision-first abstains — independent of the real corpus.
* **Real-corpus pins** (skipped when the corpus is absent): the four gold-100 cross-aggregation targets
  idx13 / idx26 / idx31 / idx46 that were棄権 before this tool existed.
"""
from __future__ import annotations

from pathlib import Path

import docx
import pytest

from src.rag import corpus
from src.rag.extract.glossary import Glossary
import importlib

from src.rag.tools import contract as C

# The package __init__ re-exports the ``corpus_aggregate`` *function*, shadowing the submodule
# attribute; import the module explicitly so ``ca`` is the module (not the function).
ca = importlib.import_module("src.rag.tools.corpus_aggregate")

_CORPUS_PRESENT = bool(corpus.walk())
_needs_corpus = pytest.mark.skipif(not _CORPUS_PRESENT, reason="corpus not present")


# --------------------------------------------------------------------------- synthetic mini-corpus
def _write_contract(path: Path, *, total_label: str, total: str, deposit_row: list[str] | None,
                    period: str | None, staff: dict[str, str]) -> None:
    """Emit a docx contract whose extracted text matches the real-corpus field formats."""
    path.parent.mkdir(parents=True, exist_ok=True)
    d = docx.Document()
    d.add_paragraph("データ分析業務委託契約書")
    d.add_paragraph(f"{total_label}：{total}円")
    if period:
        d.add_paragraph(f"本契約の契約期間は、{period}までとする。")
    for label, name in staff.items():
        d.add_paragraph(f"{label}：{name}")
    if deposit_row is not None:
        # 着手金 lives in a table row: "1 | 着手金 | 40% | 税抜 | 税額 | 税込 | 期限" (税込 = 最大)
        t = d.add_table(rows=1, cols=len(deposit_row))
        for cell, text in zip(t.rows[0].cells, deposit_row):
            cell.text = text
    d.save(str(path))


def _write_train(path: Path, rows: int) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    body = "\n".join(f"{i},{i * 2}" for i in range(1, rows + 1))
    path.write_text("id,x\n" + body + "\n", encoding="utf-8")


@pytest.fixture()
def mini(tmp_path: Path) -> tuple[Path, Glossary]:
    """A 3-project mini drive: two 固定金額契約 (着手金あり) + one time-and-materials案件."""
    root = tmp_path / "share_drive"
    proj = root / "プロジェクト"

    # ALPHA — fixed-price, 着手金 3,300,000(税込), 契約金額 6,600,000, period overlaps & >40日, 7 rows.
    _write_contract(
        proj / "アルファ株式会社" / "01.契約" / "契約書.docx",
        total_label="契約金額（税込）", total="6,600,000",
        deposit_row=["1", "着手金", "40%", "3,000,000円", "300,000円", "3,300,000円", "締結後5営業日"],
        period="2025-08-01から2025-09-20",
        staff={"エグゼクティブスポンサー": "山田 太郎", "プロジェクトマネージャー": "佐藤 花子",
               "データエンジニア": "田中 一郎"})
    _write_train(proj / "アルファ株式会社" / "03.データ" / "train.csv", 7)

    # BETA — fixed-price, 着手金 2,200,000, 契約金額 4,400,000, period no overlap, 200 rows.
    _write_contract(
        proj / "ベータ株式会社" / "01.契約" / "契約書.docx",
        total_label="契約金額（税込）", total="4,400,000",
        deposit_row=["1", "着手金", "50%", "2,000,000円", "200,000円", "2,200,000円", "締結後5営業日"],
        period="2025-01-05から2025-03-01",
        staff={"プロジェクトマネージャー": "佐藤 花子", "データエンジニア": "田中 一郎"})
    _write_train(proj / "ベータ株式会社" / "03.データ" / "train.csv", 200)

    # GAMMA — time-and-materials (no 着手金, not fixed), 想定金額 9,900,000, 有効期間 <40日, 300 rows.
    _write_contract(
        proj / "ガンマ株式会社" / "01.契約" / "契約書.docx",
        total_label="想定金額（税込）", total="9,900,000",
        deposit_row=None,
        period=None,
        staff={"エグゼクティブスポンサー": "田中 一郎"})
    # 有効期間 (別ラベル) を段落で追加 → 期間抽出が契約期間/有効期間両対応であることも確認
    gdoc = proj / "ガンマ株式会社" / "01.契約" / "契約書.docx"
    d = docx.Document(str(gdoc))
    d.add_paragraph("本契約の有効期間は、2025-08-10から2025-08-25までとする。")
    d.add_paragraph("本契約の料金体系は、time_and_materials とする。")
    d.save(str(gdoc))
    _write_train(proj / "ガンマ株式会社" / "03.データ" / "train.csv", 300)

    g = Glossary()
    g.company_to_code = {"アルファ株式会社": "ALPHA", "ベータ株式会社": "BETA", "ガンマ株式会社": "GAMMA"}
    return root, g


def _ca(mini, metric, op="max", **kw):
    root, g = mini
    return ca.corpus_aggregate(metric, op=op, corpus_dir=root, glossary=g, **kw)


def test_collect_extracts_all_fields_and_abbrev(mini):
    root, g = mini
    recs = {r.abbrev: r for r in ca.collect_contracts(corpus_dir=root, glossary=g)}
    assert set(recs) == {"ALPHA", "BETA", "GAMMA"}                 # 主略称 normalization
    assert recs["ALPHA"].contract_amount == 6_600_000
    assert recs["ALPHA"].deposit == 3_300_000                     # 税込 = max of the 着手金 row
    assert recs["ALPHA"].train_rows == 7
    assert recs["ALPHA"].period == ("2025-08-01", "2025-09-20")
    assert recs["ALPHA"].fixed is True
    assert recs["ALPHA"].staff["ES"] == "山田 太郎"
    # GAMMA is T&M: no 着手金, not fixed, 有効期間 label still parsed
    assert recs["GAMMA"].deposit is None and recs["GAMMA"].fixed is False
    assert recs["GAMMA"].period == ("2025-08-10", "2025-08-25")


def test_deposit_max_and_min(mini):
    hi = _ca(mini, "deposit", "max")
    assert C.is_contract(hi) and hi["value"]["abbrev"] == "ALPHA" and hi["value"]["value"] == 3_300_000
    assert hi["value"]["staff"]["ES"] == "山田 太郎"               # staff同梱で後続の内線引きに使える
    lo = _ca(mini, "deposit", "min")
    assert lo["value"]["abbrev"] == "BETA"                        # GAMMA(None)は候補から除外


def test_amount_per_row_fixed_only_and_round_up(mini):
    # per-row: ALPHA 6,600,000/7≈942857.14, BETA 4,400,000/200=22000 → ALPHA max。GAMMA(非固定)は除外。
    r = _ca(mini, "amount_per_row", "max", fixed_only=True)
    assert r["value"]["abbrev"] == "ALPHA"
    assert r["value"]["value"] == pytest.approx(942857.14, abs=0.01)
    up = _ca(mini, "amount_per_row", "max", fixed_only=True, round_up=True)
    assert up["value"]["value"] == 942858                         # 円単位で切り上げ
    assert up["value"]["amount"] == 6_600_000 and up["value"]["rows"] == 7


def test_staff_cross_count(mini):
    r = _ca(mini, "staff", "count")
    # 田中 一郎: ALPHA(DE)+BETA(DE)+GAMMA(ES)=3 が最頻
    assert r["value"]["top"] == "田中 一郎" and r["value"]["count"] == 3
    assert r["value"]["counts"]["佐藤 花子"] == 2
    assert r["evidence"]["tie"] is None


def test_period_overlap_filter(mini):
    r = _ca(mini, "period_days", "filter",
            overlap_start="2025-08-15", overlap_end="2025-09-07", min_days=40)
    # ALPHA (08-01..09-20, 50日, overlap) のみ。BETA=非overlap, GAMMA=15日で>40不成立。
    assert r["value"] == ["ALPHA"]
    assert r["evidence"]["matches"][0]["days"] == 50


def test_precision_guards(mini):
    assert ca.corpus_aggregate("nope")["value"] is None
    assert "unknown metric" in ca.corpus_aggregate("nope")["evidence"]["error"]
    assert ca.corpus_aggregate("deposit", "frobnicate")["value"] is None
    # window that overlaps nothing → empty list (no fabrication)
    empty = _ca(mini, "period_days", "filter",
                overlap_start="2030-01-01", overlap_end="2030-02-01", min_days=40)
    assert empty["value"] == []


def test_contract_shape_everywhere(mini):
    for op_kw in (("deposit", {}), ("staff", {"op": "count"}),
                  ("period_days", {"op": "filter"}), ("train_rows", {"op": "max"})):
        metric, kw = op_kw
        assert C.is_contract(_ca(mini, metric, **kw))


# --------------------------------------------------------------------------- real-corpus gold pins
@_needs_corpus
def test_idx13_most_involved_person():
    # データアステル社で最も多く案件に関わる人 = 斎藤 悠斗 (内線は seating_chart で 7104)
    r = ca.corpus_aggregate("staff", "count")
    assert r["value"]["top"] == "斎藤 悠斗" and r["value"]["count"] >= 4


@_needs_corpus
def test_idx26_period_overlap_over_40_days():
    r = ca.corpus_aggregate("period_days", "filter",
                            overlap_start="2025-08-15", overlap_end="2025-09-07", min_days=40)
    assert r["value"] == ["TOTO", "AOMINE"]              # gold: TOTO、AOMINE(順序含む)


@_needs_corpus
def test_idx31_fixed_price_max_amount_per_row():
    r = ca.corpus_aggregate("amount_per_row", "max", fixed_only=True, round_up=True)
    assert r["value"]["abbrev"] == "MINAMINO" and r["value"]["value"] == 1320   # gold: MINAMINO、1320円


@_needs_corpus
def test_idx46_max_deposit_case_and_its_es():
    r = ca.corpus_aggregate("deposit", "max")
    # 着手金最大 = 白峰(SHR)、その ES = 中村 誠 (内線は seating_chart で 7201)
    assert r["value"]["abbrev"] == "SHR"
    assert r["value"]["staff"]["ES"] == "中村 誠"


@_needs_corpus
def test_idx86_closed_staff_population_across_four_document_types():
    r = ca.corpus_aggregate("staff_population", "count")
    assert r["value"]["count"] == 19
    assert len(r["value"]["people"]) == 19
    assert len(r["evidence"]["selected_files"]) == 40  # 10 projects × PP/contract/PLAN/FR
    assert r["evidence"]["missing"] == []
    assert r["evidence"]["unreadable_files"] == []
    assert r["evidence"]["closure"] == {
        "authoritative_population_resolved": True,
        "inclusion_exclusion_recorded": True,
        "second_path_novel_candidates": [],
        "enumeration_count": 19,
        "aggregate_count": 19,
    }
