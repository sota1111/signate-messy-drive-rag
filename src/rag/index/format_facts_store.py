"""複合書式ファクトストア（SOT-2703 / cycle10, idx16/71）.

Sonnet gold100 cycle10 の abstain のうち idx16（青葉与信 中間報告資料の **黄色ハイライト×赤字** →
gold ``0.589``）と idx71（青嶺不動産 会議録の **太字∧下線∧イタリック** → gold ``4,250,000円``）は、
証拠が原 docx の run 書式に実在するのに **serve が到達できなかった**（SOT-2701 cycle10 分析で機械確認）:

* idx16: 『中間報告資料』がファイル名に一致せず reach 不能（社内用語集: 中間報告=M02 / 報告資料=RP）で、
  かつ ``src/rag/extract/office.py::extract_docx`` は太字・ハイライトの **単独** マーカーしか出さず、
  「黄∧赤」の複合述語を決定論に引く手段が無かった。同案件には 報告資料 が 2 本あり **どちらも** 黄×赤 run を
  持つ（04-29=0.589 / 04-09=4,620,000）ため、本文自己同定（『中間報告』を含む 報告資料 = 中間報告資料）で
  一意化しないと曖昧になる。
* idx71: serve が find_files で 会議録 を発見できず書式ツールに至らなかった。太字∧下線∧イタリックの
  複合述語も単独マーカーでは引けない。

本モジュールは build 時に一度だけ、**探索に依存せず** 全案件の全 docx を走査し、書式付き run を全数
``{project, rel, doc_kind, locator, text, attrs:{highlight, font_color, bold, underline, italic}}`` で焼く
（質問も gold も見ない網羅計算のみ）。doc-kind はパスのフォルダ名（会議録/報告資料/…）＋ 本文自己同定
（『中間報告』/『最終報告』を含む報告文書）で解決する。読み出し/配線は :mod:`src.rag.agent.format_facts_lane`。

Design invariants（sibling の format_series_store / report_attr_store / visual_store と同一）
-----------------------------------------------------------------------------------------
* **Opt-in at serve time.** :func:`enabled` (``RAG_FORMAT_FACTS``) が runtime 参照のみ gate する。
  default OFF ⇒ champion serve path は byte-identical（新キーも焼かれず build も no-op で退出しない=
  ストア自体は常に焼けるが serve は OFF なら読まない）。
* **Build は LLM フリー・追加的.** python-docx / OOXML の決定論読取のみ（``format_events`` の run 書式
  ヘルパ＋python-docx の bold/underline/italic）。暗号化 docx は passwords 経由で復号、読めない 1 本は
  スキップして継続。build で新たな genai 呼び出しはしない。
* **Question-independent.** universe は全案件の全 docx。質問も gold も参照せず、全書式 run を網羅記録する。
* **No hardcoding.** gold 値・idx 番号を一切持たない。全値は原ファイルの run テキストを逐語で読む。
* **Fail-open.** artifact 欠落・解析不能はすべて空へフォールバック（回帰ゼロ）。
"""
from __future__ import annotations

import io
import json
import os
import unicodedata
from pathlib import Path
from typing import Any, Mapping, Sequence

from config import settings
from src.rag.corpus import FileRef, nfc, walk

SCHEMA = "format-facts-store"
SCHEMA_VERSION = 2  # v2: emph_spans（docx/pptx/pdf の merged 強調span）を追加（SOT-2711 idx3/11）

_ON = {"1", "true", "yes", "on"}

# 色 run（highlight/font_color）を持ちうるのは docx のみ（``runs`` フィールドは docx 限定 = idx16 系）。
_DOC_EXTS = {"docx"}
# 強調span（太字/下線/イタリック）の抽出対象。SOT-2711: pptx/pdf の書式も font_emphasis で網羅する
# （idx11 = 青嶺 報告資料 PDF の B∧U∧I / 契約書 pptx 等）。emph_spans は merged span = 逐語列挙の単位。
_EMPH_EXTS = {"docx", "pptx", "pdf"}

# doc-kind をパスのフォルダ名/ファイル名から決めるための語彙（質問が使う自然表現）。
_KIND_FOLDERS = (
    "会議録", "議事録", "報告資料", "報告書", "提案書", "契約書", "見積書", "仕様書",
    "定例", "月次",
)
# 本文自己同定で付与する doc-kind（『中間報告』を含む報告文書 = 中間報告資料, 等）。
_SELFID_PHRASES = ("中間報告", "最終報告")


def enabled() -> bool:
    """True when the serve path may consult the composite-format facts store (default OFF — opt-in)."""
    return os.getenv("RAG_FORMAT_FACTS", "0").strip().lower() in _ON


def default_out_path() -> Path:
    return settings.ARTIFACTS_DIR / "format_facts_store.jsonl"


def default_report_path() -> Path:
    return settings.ARTIFACTS_DIR / "format_facts_store_build_report.json"


# --------------------------------------------------------------------------- helpers
def norm(text: Any) -> str:
    return unicodedata.normalize("NFKC", str(text or "")).replace(" ", "").replace("　", "").lower()


def _office_bytes(ref: FileRef) -> bytes | None:
    """暗号化 docx を passwords 経由で復号したバイト列（非暗号化なら None ⇒ path 直読）。"""
    try:
        from src.rag.tools.highlight_extract import _office_bytes as _ob
        return _ob(ref, None)
    except Exception:  # noqa: BLE001 — 復号不能は呼び出し側で path 直読/スキップ
        return None


def _open_docx(ref: FileRef):
    data = _office_bytes(ref)
    import docx
    return docx.Document(io.BytesIO(data) if data else str(ref.path))


# --------------------------------------------------------------------------- run 書式抽出
def _run_attrs(run, rpr) -> dict[str, Any]:
    """1 run の実効書式 {highlight, font_color, bold, underline, italic}（未設定は None/False）。"""
    from src.rag.tools.font_emphasis import _as_bool
    from src.rag.tools.format_events import _resolve_font_color, _run_highlight_name

    font_name, _hex, _theme = _resolve_font_color(rpr.find(_W_COLOR) if rpr is not None else None)
    highlight = _run_highlight_name(rpr)
    font = getattr(run, "font", None)
    bold = _as_bool(run.bold if run.bold is not None else (font.bold if font is not None else None))
    underline = _as_bool(
        run.underline if run.underline is not None else (font.underline if font is not None else None))
    italic = _as_bool(run.italic if run.italic is not None else (font.italic if font is not None else None))
    return {"highlight": highlight, "font_color": font_name,
            "bold": bold, "underline": underline, "italic": italic}


_W_COLOR = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color"


def _has_any_format(attrs: Mapping[str, Any]) -> bool:
    return bool(attrs.get("highlight") or attrs.get("font_color")
               or attrs.get("bold") or attrs.get("underline") or attrs.get("italic"))


def _doc_runs(ref: FileRef) -> list[dict[str, Any]]:
    """docx 本文＋表セルの書式付き run を全数（順序保存, 書式ゼロの run は捨てる）。"""
    _W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    out: list[dict[str, Any]] = []

    def emit(paragraphs, *, loc_prefix: str) -> None:
        for pi, para in enumerate(paragraphs, 1):
            for ri, run in enumerate(para.runs, 1):
                text = (run.text or "").strip()
                if not text:
                    continue
                rpr = run._r.find(_W + "rPr")
                attrs = _run_attrs(run, rpr)
                if not _has_any_format(attrs):
                    continue
                out.append({"text": text, "attrs": attrs,
                            "loc": f"{loc_prefix}para:{pi}:run:{ri}"})

    try:
        d = _open_docx(ref)
    except Exception:  # noqa: BLE001 — 読めない/復号不能 docx はスキップ
        return out
    emit(d.paragraphs, loc_prefix="")
    for ti, table in enumerate(d.tables, 1):
        for rr, trow in enumerate(table.rows, 1):
            for cc, tcell in enumerate(trow.cells, 1):
                emit(tcell.paragraphs, loc_prefix=f"表{ti}:R{rr}C{cc}:")
    return out


# --------------------------------------------------------------------------- 強調span（merged, docx/pptx/pdf）
def _span_loc(ev: Mapping[str, Any]) -> str:
    """font_emphasis の evidence を人間可読な locator 文字列へ（paragraph/slide/shape/page/bbox）。"""
    parts: list[str] = []
    for key in ("sheet", "cell", "paragraph", "slide", "shape", "page"):
        v = ev.get(key)
        if v not in (None, ""):
            parts.append(f"{key}:{v}")
    bbox = ev.get("bbox")
    if bbox:
        parts.append("bbox:" + ",".join(str(x) for x in bbox))
    return "/".join(parts)


def _emph_spans(ref: FileRef) -> list[dict[str, Any]]:
    """全書式強調 span を font_emphasis で **merged** 抽出（bold/underline/italic, 文書順）.

    ``font_emphasis`` は隣接する同一装飾 run/glyph を 1 span に束ねる（例: docx で「30」「分単位」に
    分割された太字 run を ``30分単位`` に、pdf の sheared「4,675,000円」を 1 glyph_run に）。この merged
    形が「太字箇所をすべて列挙」型（idx3）や PDF の B∧U∧I 一意束縛（idx11）の逐語単位になる。装飾ゼロは
    font_emphasis が返さない。読めない/未対応拡張子は空へ（fail-open）。
    """
    try:
        from src.rag.tools.font_emphasis import font_emphasis
        res = font_emphasis(ref)  # require=None ⇒ 装飾 ≥1 の全 span
    except Exception:  # noqa: BLE001 — 復号不能/解析不能/未対応は空へ（回帰ゼロ）
        return []
    out: list[dict[str, Any]] = []
    for item in res.get("value", []):
        text = str(item.get("value", "")).strip()
        if not text:
            continue
        method = item.get("method", {}) or {}
        attrs = {"bold": bool(method.get("bold")),
                 "underline": bool(method.get("underline")),
                 "italic": bool(method.get("italic"))}
        out.append({"text": text, "attrs": attrs, "loc": _span_loc(item.get("evidence", {}) or {})})
    return out


# --------------------------------------------------------------------------- doc-kind
def _folder_kinds(rel: str, name: str) -> list[str]:
    """パスのフォルダ名/ファイル名から doc-kind トークン（会議録/報告資料/…）を決定論抽出。"""
    hay = nfc(rel) + " " + nfc(name)
    kinds: list[str] = []
    for k in _KIND_FOLDERS:
        if k in hay and k not in kinds:
            kinds.append(k)
    return kinds


def _selfid_kinds(body: str, folder_kinds: Sequence[str]) -> list[str]:
    """本文自己同定 doc-kind（『中間報告』/『最終報告』を含む 報告資料/報告書 = 中間報告資料 等）。"""
    # 報告系文書だけ自己同定（会議録/契約書に『中間報告』が言及されても doc-kind にはしない）。
    is_report = any(k in ("報告資料", "報告書") for k in folder_kinds)
    if not is_report:
        return []
    out: list[str] = []
    for phrase in _SELFID_PHRASES:
        if phrase in body and phrase not in out:
            out.append(phrase)
    return out


def _doc_body(ref: FileRef) -> str:
    """docx 本文プレーンテキスト（自己同定用。genai 非使用, 復号対応, 失敗は空）。"""
    try:
        d = _open_docx(ref)
    except Exception:  # noqa: BLE001
        return ""
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for t in d.tables:
        for r in t.rows:
            parts.append(" | ".join(c.text for c in r.cells))
    return nfc("\n".join(parts))


# --------------------------------------------------------------------------- universe / build
def _universe(refs: Sequence[FileRef]) -> list[FileRef]:
    out: list[FileRef] = []
    for r in refs:
        if r.ext not in _EMPH_EXTS:
            continue
        name = nfc(r.name).lower()
        rel = nfc(r.rel).lower()
        if name.startswith("~$") or "/old/" in rel or name.startswith("old"):
            continue
        out.append(r)
    out.sort(key=lambda r: nfc(r.rel))
    return out


def compute_doc(ref: FileRef) -> dict[str, Any] | None:
    """1 文書の書式事実を組む（docx の色 run + 全拡張子の merged 強調span + doc-kind）.

    ``runs`` は docx 限定の色付き run（highlight/font_color/太字…= idx16/71 の単発 lookup 用）、
    ``emph_spans`` は docx/pptx/pdf の merged 強調 span（太字/下線/イタリック = idx3 列挙 / idx11 PDF）。
    両方空なら None（書式証拠のない文書は欠測を偽装せず記録しない）。
    """
    runs = _doc_runs(ref) if ref.ext in _DOC_EXTS else []
    emph = _emph_spans(ref)
    if not runs and not emph:
        return None
    folder_kinds = _folder_kinds(ref.rel, ref.name)
    # 本文自己同定（『中間報告』等）は docx 本文からのみ（pptx/pdf は folder-kind で足りる）。
    selfid = (_selfid_kinds(_doc_body(ref), folder_kinds)
              if folder_kinds and ref.ext in _DOC_EXTS else [])
    doc_kind = folder_kinds + [k for k in selfid if k not in folder_kinds]
    return {
        "project": nfc(ref.project), "rel": nfc(ref.rel), "name": nfc(ref.name),
        "doc_kind": doc_kind, "ext": ref.ext,
        "n_runs": len(runs), "runs": runs,
        "n_emph": len(emph), "emph_spans": emph,
    }


def build(refs: Sequence[FileRef] | None = None, *, out: Path | None = None,
          write_report: bool = True) -> dict[str, Any]:
    """全案件の全 docx を走査し、複合書式ファクトストアを構築（LLM 非使用・べき等・fail-open）。"""
    refs = list(refs) if refs is not None else list(walk())
    universe = _universe(refs)
    records: list[dict[str, Any]] = []
    skipped = 0
    for ref in universe:
        try:
            rec = compute_doc(ref)
        except Exception:  # noqa: BLE001 — 一文書の失敗で build を落とさない
            rec = None
        if rec is not None:
            records.append(rec)
        else:
            skipped += 1
    stats = write_store(records, out)
    report = {
        "schema": SCHEMA, "version": SCHEMA_VERSION,
        "universe": len(universe), "records": len(records), "skipped": skipped,
        "total_runs": sum(r["n_runs"] for r in records),
        "total_emph": sum(r.get("n_emph", 0) for r in records),
        "by_ext": {e: sum(1 for r in records if r.get("ext") == e) for e in sorted(_EMPH_EXTS)},
        "docs_with_selfid": sum(1 for r in records
                                if any(k in _SELFID_PHRASES for k in r.get("doc_kind", []))),
    }
    if write_report:
        rp = default_report_path()
        rp.parent.mkdir(parents=True, exist_ok=True)
        rp.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"records": len(records), "docs": stats["docs"], "report": report}


# --------------------------------------------------------------------------- io
def write_store(records: Sequence[Mapping[str, Any]], path: Path | None = None) -> dict[str, int]:
    """Atomically write a reproducible JSONL store (schema header + rel-sorted rows)."""
    out = Path(path) if path is not None else default_out_path()
    out.parent.mkdir(parents=True, exist_ok=True)
    ordered = sorted(records, key=lambda r: (r.get("project", ""), r.get("rel", "")))
    tmp = out.with_suffix(out.suffix + ".tmp")
    with tmp.open("w", encoding="utf-8", newline="\n") as handle:
        handle.write(json.dumps({"schema": SCHEMA, "version": SCHEMA_VERSION},
                                ensure_ascii=False, sort_keys=True) + "\n")
        for rec in ordered:
            handle.write(json.dumps(rec, ensure_ascii=False, sort_keys=True) + "\n")
    tmp.replace(out)
    reset_cache()
    return {"docs": len(ordered)}


_LOAD_CACHE: dict[str, list[dict[str, Any]]] = {}


def reset_cache() -> None:
    _LOAD_CACHE.clear()


def load(path: Path | None = None) -> list[dict[str, Any]]:
    """Load the format-facts rows (memoized). ``[]`` when absent/unreadable/schema-mismatch (回帰ゼロ)."""
    out = Path(path) if path is not None else default_out_path()
    key = str(out)
    if key in _LOAD_CACHE:
        return _LOAD_CACHE[key]
    rows: list[dict[str, Any]] = []
    try:
        with open(out, encoding="utf-8") as handle:
            header = handle.readline()
            meta = json.loads(header) if header.strip() else {}
            if isinstance(meta, dict) and meta.get("schema") == SCHEMA \
                    and meta.get("version") == SCHEMA_VERSION:
                for line in handle:
                    line = line.strip()
                    if line:
                        rows.append(json.loads(line))
    except Exception:  # noqa: BLE001
        rows = []
    _LOAD_CACHE[key] = rows
    return rows


if __name__ == "__main__":
    summary = build()
    print(f"[build] format_facts_store records={summary['records']} "
          f"docs={summary['docs']} total_runs={summary['report']['total_runs']} "
          f"total_emph={summary['report']['total_emph']} by_ext={summary['report']['by_ext']} "
          f"-> {default_out_path()}")
