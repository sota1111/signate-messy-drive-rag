"""文書内数値の式適用ストア（SOT-2694 / cycle8 C5, idx68/50）.

Sonnet gold100 cycle6/7/8 の abstain idx68/50 は、答えに必要な数値がコーパス内テキストに **到達可能**
（証拠は既存ストアに焼き込み済み）でありながら、投資者ループが予算内に「記載式への数値代入」や
「入れ子表からの統計量差分」を導けず棄権していた（``docs/ai/sonnet_cycle_analysis/cycle8.md`` C5）。
欠けているのは証拠ではなく **決定論的な式適用/差分の浅到達レーン** である。

本モジュールは build 時に一度だけ、質問も gold も見ずに全数走査し、次の 2 種のレコードを焼く（LLM 非使用）:

* **formula レコード（idx68 型）** — 画像 OCR ストア（:mod:`src.rag.index.image_ocr_store`, 既存成果）の各
  ページ本文に「``<名称>＝<式>``」の記載式があり、その式の変数名（``生産性向上率`` 等）が同じページの
  ラベル付き数値（``+22.6%`` / ``3.7倍`` 等）へ一意に束縛できるとき、Decimal で式を評価し値を焼く。
  例: ``投資実装係数＝（生産性向上率＋コスト削減率）×ROI倍率`` に ``(0.226+0.152)×3.7`` を代入 ⇒ ``1.3986``。
* **stat_table レコード（idx50 型）** — docx の **入れ子表**（``python-docx`` の ``document.tables`` は
  トップレベル表しか返さず、セル内の入れ子表を落とすため native 抽出が到達できない）を再帰的に走査し、
  ヘッダに統計量ラベル（``中央値`` / ``上位90%`` / ``下位10%`` 等）を持つ「行キー×統計量列」表を焼く。
  例: ``Salary.com`` 行の ``上位90%``(137,000) と ``中央値``(123,778) の差 ⇒ ``13,222``。

読み出し/配線は :mod:`src.rag.agent.formula_apply_lane`。

Design invariants（sibling の rate_table_store / case_finance_store と同一）:
* **Opt-in at serve time.** :func:`enabled` (``RAG_FORMULA_APPLY``) が runtime 参照のみ gate する。
  default OFF ⇒ champion serve path は byte-identical。
* **Build は LLM フリー・追加的.** image_ocr ストア(既存)と python-docx を読むだけ。読めない文書は 1 件
  スキップして継続。
* **Question-independent / No hardcoding.** universe は全 docx／全 OCR ページ。質問も gold も idx も
  参照せず、構造（記載式・統計量表）を検出して網羅計算する。全値は原ファイル由来で出典（doc/ページ）付き。
* **Fail-open.** artifact 欠落・解析不能はすべて空へフォールバック（回帰ゼロ）。
"""
from __future__ import annotations

import json
import os
import re
import unicodedata
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any, Iterator, Mapping, Sequence

from config import settings
from src.rag.corpus import FileRef, nfc, walk
from src.rag.index import image_ocr_store as _image_ocr

SCHEMA = "formula-apply-store"
SCHEMA_VERSION = 1

_ON = {"1", "true", "yes", "on"}


def enabled() -> bool:
    """True when the serve path may consult the formula-apply store (default OFF — opt-in)."""
    return os.getenv("RAG_FORMULA_APPLY", "0").strip().lower() in _ON


def default_out_path() -> Path:
    return settings.ARTIFACTS_DIR / "formula_apply_store.jsonl"


def default_report_path() -> Path:
    return settings.ARTIFACTS_DIR / "formula_apply_store_build_report.json"


# --------------------------------------------------------------------------- normalization helpers
def norm(text: Any) -> str:
    """NFKC + strip all whitespace + casefold（表記ゆれ吸収）。"""
    return re.sub(r"\s+", "", unicodedata.normalize("NFKC", str(text or ""))).casefold()


def _nfkc(text: Any) -> str:
    return unicodedata.normalize("NFKC", str(text or ""))


def _to_number(text: str) -> "int | float | None":
    """セル/トークンの先頭数値を取り出す（``123,778`` → 123778、``1.5`` → 1.5）。無ければ None。"""
    m = re.search(r"[+-]?\d[\d,]*(?:\.\d+)?", _nfkc(text))
    if not m:
        return None
    raw = m.group(0).replace(",", "").lstrip("+")
    try:
        if "." in raw:
            return float(raw)
        return int(raw)
    except ValueError:
        return None


# =========================================================================== formula records (idx68)
# 記載式「<名称>＝<式>」の名称は、係数/指数/スコア等の派生量であることを要求（過検出抑止）。
_NAME_SUFFIXES = ("係数", "指数", "スコア", "比率", "指標")
# 変数名の接尾辞 → 期待単位（束縛の曖昧性を切る）。
#   ``…倍率`` / ``…倍`` は 倍 値、``…率`` は % 値を要求する。
_PCT_RE = re.compile(r"([+-]?\d[\d,]*(?:\.\d+)?)\s*(?:%|％|パーセント)")
_MULT_RE = re.compile(r"([+-]?\d[\d,]*(?:\.\d+)?)\s*倍")
# 式トークンの区切り（演算子・括弧・空白）。
_OP_SPLIT = re.compile(r"[()（）+\-*/×÷·・、,\s]+")
_OPERATORS = set("+-*/×÷")


def _var_base_unit(var: str) -> "tuple[str, str | None]":
    """変数名 → (束縛に使う基底ラベル, 期待単位 '%'|'倍'|None)。"""
    v = _nfkc(var)
    if v.endswith("倍率"):
        return v[:-2], "倍"
    if v.endswith("倍"):
        return v[:-1], "倍"
    if v.endswith("率"):
        return v[:-1], "%"
    return v, None


def _is_number_token(tok: str) -> bool:
    return bool(re.fullmatch(r"[+-]?\d[\d,]*(?:\.\d+)?", _nfkc(tok).strip()))


def _slash_split(line: str) -> list[str]:
    parts = re.split(r"[／/]", _nfkc(line))
    return [p.strip() for p in parts if p.strip()]


def _slashzip_bindings(lines: list[str]) -> dict[str, tuple[str, str]]:
    """隣接する「値の /-並び」と「ラベルの /-並び」を同数なら位置対応で束縛。

    例: ``+15.2% / +22.6%`` と ``コスト削減 / 生産性向上`` ⇒
    ``{コスト削減: ('15.2','%'), 生産性向上: ('22.6','%')}``（正規化キー）。
    """
    out: dict[str, tuple[str, str]] = {}
    for i in range(len(lines) - 1):
        a, b = _slash_split(lines[i]), _slash_split(lines[i + 1])
        if len(a) < 2 or len(a) != len(b):
            continue
        a_num = [t for t in a if _num_unit(t) is not None]
        b_num = [t for t in b if _num_unit(t) is not None]
        # 一方が全て数値+単位、他方が全て非数値ラベルのときだけ対応づける。
        if len(a_num) == len(a) and len(b_num) == 0:
            values, labels = a, b
        elif len(b_num) == len(b) and len(a_num) == 0:
            values, labels = b, a
        else:
            continue
        for label, value in zip(labels, values):
            nu = _num_unit(value)
            if nu is not None:
                out[norm(label)] = nu
    return out


def _num_unit(token: str) -> "tuple[str, str] | None":
    """トークンから (数値文字列, 単位 '%'|'倍') を取り出す。%・倍 以外は None。"""
    t = _nfkc(token)
    m = _PCT_RE.search(t)
    if m:
        return m.group(1).replace(",", ""), "%"
    m = _MULT_RE.search(t)
    if m:
        return m.group(1).replace(",", ""), "倍"
    return None


def _bind_variable(base: str, want_unit: "str | None", lines: list[str],
                   slashzip: dict[str, tuple[str, str]]) -> "Decimal | None":
    """基底ラベル base を、ページ内の一意なラベル付き数値へ束縛して Decimal を返す（曖昧なら None）。"""
    key = norm(base)
    # (1) 並記リストからの束縛（単位一致のみ採用）。
    hit = slashzip.get(key)
    if hit is not None and (want_unit is None or hit[1] == want_unit):
        return _decimal_of(hit[0], hit[1])
    # (2) ラベル出現行±1 の近傍にある「期待単位の数値」が一意なら束縛。
    cands: list[tuple[str, str]] = []
    for idx, line in enumerate(lines):
        if key and key in norm(line):
            for j in (idx - 1, idx, idx + 1):
                if 0 <= j < len(lines):
                    nu = _num_unit(lines[j])
                    if nu is not None and (want_unit is None or nu[1] == want_unit):
                        cands.append(nu)
    uniq = {c for c in cands}
    if len(uniq) == 1:
        nu = next(iter(uniq))
        return _decimal_of(nu[0], nu[1])
    # (3) 期待単位が明示され、その単位の数値がページ全体で唯一なら束縛（近傍で拾えない配置の保険）。
    if want_unit is not None:
        page_all: list[tuple[str, str]] = []
        for line in lines:
            for m in _PCT_RE.finditer(_nfkc(line)):
                page_all.append((m.group(1).replace(",", ""), "%"))
            for m in _MULT_RE.finditer(_nfkc(line)):
                page_all.append((m.group(1).replace(",", ""), "倍"))
        same = [nu for nu in page_all if nu[1] == want_unit]
        if len({s for s in same}) == 1:
            return _decimal_of(same[0][0], same[0][1])
    return None


def _decimal_of(numstr: str, unit: str) -> "Decimal | None":
    try:
        n = Decimal(str(numstr).lstrip("+"))
    except (InvalidOperation, ValueError):
        return None
    if unit == "%":
        return n / Decimal(100)
    return n  # 倍 / none → そのまま


def _eval_expr(expr: str, env: Mapping[str, Decimal]) -> "Decimal | None":
    """変数を Decimal 値へ置換した式を安全に評価（+ - * / と括弧のみ、eval 不使用）。"""
    # 変数名を長い順に Decimal リテラルへ置換（部分一致の取り違え防止）。
    substituted = _nfkc(expr).replace("×", "*").replace("÷", "/").replace("（", "(").replace("）", ")")
    for name in sorted(env, key=len, reverse=True):
        substituted = substituted.replace(name, f"({env[name]})")
    # 置換後に許可文字（数字・演算子・括弧・小数点・空白）以外が残れば未束縛変数 ⇒ 評価不能。
    if re.search(r"[^0-9.+\-*/()\s]", substituted):
        return None
    try:
        return _eval_tokens(_tokenize(substituted))
    except (ArithmeticError, ValueError, IndexError):
        return None


def _tokenize(s: str) -> list[str]:
    return re.findall(r"\d+(?:\.\d+)?|[()+\-*/]", s)


def _eval_tokens(tokens: list[str]) -> Decimal:
    """再帰下降で + - * / と括弧を Decimal 評価する小さな安全パーサ。"""
    pos = 0

    def parse_expr() -> Decimal:
        nonlocal pos
        value = parse_term()
        while pos < len(tokens) and tokens[pos] in ("+", "-"):
            op = tokens[pos]; pos += 1
            rhs = parse_term()
            value = value + rhs if op == "+" else value - rhs
        return value

    def parse_term() -> Decimal:
        nonlocal pos
        value = parse_factor()
        while pos < len(tokens) and tokens[pos] in ("*", "/"):
            op = tokens[pos]; pos += 1
            rhs = parse_factor()
            value = value * rhs if op == "*" else value / rhs
        return value

    def parse_factor() -> Decimal:
        nonlocal pos
        tok = tokens[pos]
        if tok == "(":
            pos += 1
            value = parse_expr()
            if pos >= len(tokens) or tokens[pos] != ")":
                raise ValueError("unbalanced")
            pos += 1
            return value
        if tok in ("+", "-"):
            pos += 1
            v = parse_factor()
            return v if tok == "+" else -v
        pos += 1
        return Decimal(tok)

    result = parse_expr()
    if pos != len(tokens):
        raise ValueError("trailing tokens")
    return result


def _formula_records_for_page(ref: FileRef, locus: str, text: str) -> list[dict[str, Any]]:
    """1 ページ本文から記載式レコードを抽出（束縛が完全なものだけ）。"""
    lines = [ln.strip() for ln in _nfkc(text).splitlines() if ln.strip()]
    slashzip = _slashzip_bindings(lines)
    out: list[dict[str, Any]] = []
    for line in lines:
        if "=" not in line:
            continue
        for chunk in re.split(r"[。\n]", line):
            if "=" not in chunk:
                continue
            lhs, rhs = chunk.split("=", 1)
            # 名称 = 「=」直前の識別子（先頭の「注釈:」等の接頭を除く）。
            name = re.split(r"[：:\s]", lhs.strip())[-1].strip()
            if not name or not any(name.endswith(sfx) or sfx in name for sfx in _NAME_SUFFIXES):
                continue
            rhs = rhs.strip()
            raw_tokens = [t for t in _OP_SPLIT.split(rhs) if t]
            var_tokens = [t for t in raw_tokens if not _is_number_token(t)]
            if not var_tokens or not any(op in rhs for op in _OPERATORS):
                continue
            env: dict[str, Decimal] = {}
            bindings: dict[str, Any] = {}
            ok = True
            for var in var_tokens:
                base, unit = _var_base_unit(var)
                val = _bind_variable(base, unit, lines, slashzip)
                if val is None:
                    ok = False
                    break
                env[var] = val
                bindings[var] = {"base": base, "expected_unit": unit, "value": str(val)}
            if not ok:
                continue
            value = _eval_expr(rhs, env)
            if value is None:
                continue
            out.append({
                "kind": "formula", "doc_id": nfc(ref.rel), "project": nfc(ref.project),
                "doc_name": nfc(ref.name), "locus": nfc(locus),
                "formula_name": name, "expression": rhs,
                "bindings": bindings, "value": _decimal_repr(value),
            })
    return out


def _decimal_repr(value: Decimal) -> str:
    """末尾の余分な 0 を落とした素直な小数表現（``1.3986`` / ``2`` 等）。"""
    v = value.normalize()
    # normalize() は 100 → 1E+2 のような指数表記を作るので固定小数へ戻す。
    if v == v.to_integral_value():
        return str(v.quantize(Decimal(1)))
    return format(v, "f")


def compute_formulas(refs: Sequence[FileRef]) -> list[dict[str, Any]]:
    """全 docx/pptx/pdf の画像 OCR ページ本文から記載式レコードを網羅抽出。"""
    out: list[dict[str, Any]] = []
    by_rel: dict[str, list[dict[str, Any]]] = {}
    for rec in _image_ocr.load():
        by_rel.setdefault(nfc(rec.get("rel") or ""), []).append(rec)
    for ref in refs:
        for rec in by_rel.get(nfc(ref.rel), []):
            text = rec.get("full_text") or ""
            if "=" not in _nfkc(text):
                continue
            try:
                out.extend(_formula_records_for_page(ref, str(rec.get("locus") or ""), text))
            except Exception:  # noqa: BLE001 — 1 ページの解析失敗が build を壊さない
                continue
    return out


# ======================================================================= stat_table records (idx50)
# ヘッダのセルが統計量列であることを示す語（部分一致）。行キー列/メタ列と区別する。
_STAT_CUES = ("中央値", "中位数", "平均値", "平均", "上位90", "上位10", "上位25", "上位75",
              "下位90", "下位10", "下位25", "下位75", "最高水準", "最低水準",
              "最大値", "最小値", "第1四分位", "第3四分位", "四分位", "パーセンタイル", "percentile")
_UNIT_HINTS = (("米ドル", "ドル"), ("ドル", "ドル"), ("円", "円"), ("%", "%"), ("％", "%"))


def _iter_docx_tables(container: Any) -> Iterator[Any]:
    """トップレベル＋入れ子（セル内）の全表を再帰的に列挙する。"""
    for tbl in container.tables:
        yield tbl
        for row in tbl.rows:
            for cell in row.cells:
                yield from _iter_docx_tables(cell)


def _stat_columns(header_cells: list[str]) -> list[int]:
    return [i for i, c in enumerate(header_cells) if any(cue in c for cue in _STAT_CUES)]


def _detect_unit(header_cells: list[str]) -> "str | None":
    joined = _nfkc(" ".join(header_cells))
    for needle, unit in _UNIT_HINTS:
        if needle in joined:
            return unit
    return None


def _table_rows(tbl: Any) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in tbl.rows:
        cells = [_nfkc(c.text).replace("\n", " ").strip() for c in row.cells]
        # 末尾の空セル（結合セルの折返し由来）を落とす。
        while cells and not cells[-1]:
            cells.pop()
        rows.append(cells)
    return rows


def _stat_table_record(ref: FileRef, tbl: Any) -> "dict[str, Any] | None":
    rows = _table_rows(tbl)
    header_idx = next((i for i, r in enumerate(rows) if len(_stat_columns(r)) >= 2), None)
    if header_idx is None:
        return None
    header = rows[header_idx]
    stat_cols = _stat_columns(header)
    if not stat_cols or 0 in stat_cols:  # 行キー列（0 列目）は統計量列であってはならない
        return None
    unit = _detect_unit(header)
    out_rows: list[dict[str, Any]] = []
    for r in rows[header_idx + 1:]:
        if not r or not r[0].strip():
            continue
        key = r[0].strip()
        if _to_number(key) is not None:  # 行キーが純粋な数値ならデータ行ではない
            continue
        cells: list[dict[str, Any]] = []
        for c in stat_cols:
            val = _to_number(r[c]) if c < len(r) else None
            cells.append({"header": header[c], "value": val})
        if any(cell["value"] is not None for cell in cells):
            out_rows.append({"key": key, "key_norm": norm(key), "cells": cells})
    if len(out_rows) < 1:
        return None
    return {
        "kind": "stat_table", "doc_id": nfc(ref.rel), "project": nfc(ref.project),
        "doc_name": nfc(ref.name),
        "header": [header[c] for c in stat_cols], "unit": unit, "rows": out_rows,
    }


def compute_stat_tables(refs: Sequence[FileRef]) -> list[dict[str, Any]]:
    """全 docx の（入れ子含む）統計量表を網羅抽出（python-docx のみ・LLM 非使用）。"""
    try:
        import docx  # noqa: F401 — availability check
    except Exception:  # noqa: BLE001
        return []
    out: list[dict[str, Any]] = []
    for ref in refs:
        if ref.ext != "docx" or nfc(ref.name).startswith("~$"):
            continue
        try:
            import docx
            document = docx.Document(str(ref.path))
        except Exception:  # noqa: BLE001 — 暗号化/破損 docx はスキップ
            continue
        try:
            for tbl in _iter_docx_tables(document):
                rec = _stat_table_record(ref, tbl)
                if rec is not None:
                    out.append(rec)
        except Exception:  # noqa: BLE001
            continue
    return out


# --------------------------------------------------------------------------- build / io
def _universe(refs: Sequence[FileRef]) -> list[FileRef]:
    out = [r for r in refs if r.ext in ("docx", "pptx", "pdf") and not nfc(r.name).startswith("~$")]
    out.sort(key=lambda r: r.rel)
    return out


def build(refs: Sequence[FileRef] | None = None, *, out: Path | None = None,
          write_report: bool = True) -> dict[str, Any]:
    """記載式適用＋統計量表差分レーンのストアを焼く（LLM 非使用・べき等・fail-open）。"""
    refs = list(refs) if refs is not None else list(walk())
    universe = _universe(refs)
    records: list[dict[str, Any]] = []
    try:
        records.extend(compute_formulas(universe))
    except Exception:  # noqa: BLE001
        pass
    try:
        records.extend(compute_stat_tables(universe))
    except Exception:  # noqa: BLE001
        pass
    stats = write_store(records, out)
    report = {
        "schema": SCHEMA, "version": SCHEMA_VERSION,
        "universe": len(universe), "records": len(records),
        "formulas": sum(1 for r in records if r.get("kind") == "formula"),
        "stat_tables": sum(1 for r in records if r.get("kind") == "stat_table"),
    }
    if write_report:
        rp = default_report_path()
        rp.parent.mkdir(parents=True, exist_ok=True)
        rp.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"records": stats["records"], "report": report}


def write_store(records: Sequence[Mapping[str, Any]], path: Path | None = None) -> dict[str, int]:
    """Atomically write a reproducible JSONL store (schema header + (kind, doc_id) sorted rows)."""
    out = Path(path) if path is not None else default_out_path()
    out.parent.mkdir(parents=True, exist_ok=True)
    ordered = sorted(records, key=lambda r: (str(r.get("kind", "")), str(r.get("doc_id", "")),
                                             str(r.get("locus", "")), str(r.get("formula_name", ""))))
    tmp = out.with_suffix(out.suffix + ".tmp")
    with tmp.open("w", encoding="utf-8", newline="\n") as handle:
        handle.write(json.dumps({"schema": SCHEMA, "version": SCHEMA_VERSION},
                                ensure_ascii=False, sort_keys=True) + "\n")
        for rec in ordered:
            handle.write(json.dumps(rec, ensure_ascii=False, sort_keys=True) + "\n")
    tmp.replace(out)
    reset_cache()
    return {"records": len(ordered)}


_LOAD_CACHE: dict[str, list[dict[str, Any]]] = {}


def reset_cache() -> None:
    _LOAD_CACHE.clear()


def load(path: Path | None = None) -> list[dict[str, Any]]:
    """Load the formula-apply records (memoized). ``[]`` when absent/unreadable/schema-mismatch (回帰ゼロ)."""
    out = Path(path) if path is not None else default_out_path()
    key = str(out)
    if key in _LOAD_CACHE:
        return _LOAD_CACHE[key]
    rows: list[dict[str, Any]] = []
    try:
        with open(out, encoding="utf-8") as handle:
            header = handle.readline()
            meta = json.loads(header) if header.strip() else {}
            if isinstance(meta, dict) and meta.get("schema") == SCHEMA \
                    and meta.get("version") == SCHEMA_VERSION:
                for line in handle:
                    line = line.strip()
                    if line:
                        rows.append(json.loads(line))
    except Exception:  # noqa: BLE001
        rows = []
    _LOAD_CACHE[key] = rows
    return rows


def records_of(kind: str, *, path: Path | None = None) -> list[dict[str, Any]]:
    return [r for r in load(path) if r.get("kind") == kind]


def docs_for(kind: str, hint: str, *, path: Path | None = None) -> list[dict[str, Any]]:
    """project / doc 名の部分一致で kind レコードを引く（空白/ケース無視）。"""
    h = norm(hint)
    rows = records_of(kind, path=path)
    if not h:
        return rows
    return [r for r in rows if h in norm(r.get("doc_id")) or h in norm(r.get("project"))
            or h in norm(r.get("doc_name"))]


if __name__ == "__main__":
    summary = build()
    rep = summary["report"]
    print(f"[build] formula_apply_store records={summary['records']} "
          f"formulas={rep['formulas']} stat_tables={rep['stat_tables']} -> {default_out_path()}")
