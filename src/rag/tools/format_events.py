"""format_events — deterministic OOXML *semantic* format extraction (SOT-2585).

Third sibling of :mod:`src.rag.tools.highlight_extract` (colour fills / markers) and
:mod:`src.rag.tools.font_emphasis` (太字/下線/イタリック). Where those surface a single formatting
property, this tool surfaces the **structured format events** whose evidence the flattened text and the
single-property tools miss entirely — turning three questions that scored UNANSWERABLE
(``PARSER_CAPABILITY_MISS``) into deterministic, filterable evidence:

* **idx16** — 「黄色ハイライトかつ赤字となっている部分」 needs a *composite* predicate
  (``background==黄 AND font.color==赤``) evaluated on a docx run's **effective style**, not one property.
* **idx49** — 「コメントがついている部分をそのまま抽出」 needs the docx **comment anchor text** (the span
  the comment is attached to), recovered from ``document.xml`` ``w:commentRangeStart/End`` ⇄
  ``comments.xml`` — python-docx 1.1.x has no high-level comment API, so the XML-direct path is primary.
* **idx65** — 「黄色ハイライトになっているセルの条件」 needs the Excel **conditional-format rule** as a
  ``FORMAT_EVENT`` (type / operator / formula / priority / stopIfTrue / dxf.fill / dxf.font / dxf.border
  / dxfId), because a rule-applied fill never touches ``cell.fill``.

Three-level extraction (deep-research §書式):

* **Level 1 direct style** — ``cell.fill`` / ``cell.font`` (xlsx), run ``w:color`` / ``w:highlight`` /
  ``w:shd`` (docx). Deterministic, always available.
* **Level 2 conditional rule semantics** — Excel ``cfRule`` + ``dxf`` surfaced as a ``FORMAT_EVENT`` with
  a human-readable condition and, for the rules we can evaluate deterministically (``cellIs``), the
  numeric cells that satisfy it (``openpyxl`` is not a calc engine, and ``data_only=True`` exposes only
  the *cached* value — so only self-contained ``cellIs`` rules are evaluated; others are surfaced as the
  rule alone).
* **Level 3 rendered fallback** — intentionally NOT here: LibreOffice/Vision rendering is
  non-deterministic and belongs to the existing vision path, never this deterministic index.

Composite predicate
-------------------
``format_events(file, fill="黄", font_color="赤")`` returns only items whose **effective** background
classifies to 黄 *and* whose **effective** font colour classifies to 赤 — the idx16 AND. Each of ``fill``
/ ``font_color`` is an independent, colour-family-aware filter (reusing the same coarse Japanese colour
vocabulary and 青⊇水色 family as ``highlight_extract``); omit both to enumerate every event.

Machine-searchable notation
---------------------------
:func:`fmt_notation` renders one event as an ``__FMT__`` line
(``__FMT__ font.color=赤 background=黄 background.source=highlight style.origin=direct``) so the extract
face and the typed evidence-index can filter formatting predicates lexically, not by embedding similarity.

Gating
------
Everything is behind ``RAG_FORMAT_EVENTS`` (default OFF), mirroring ``RAG_FONT_EMPHASIS`` /
``RAG_HIGHLIGHT_EXTRA`` / ``RAG_STRUCTURE_STORE``: with the flag unset the champion serve path (tool set,
prompt, read_office face) is byte-identical. Focused/offline verification sets it on.

Contract
--------
Returns the unified :mod:`src.rag.tools.contract` shape ``{value, evidence, method}`` where ``value`` is
a list of per-event contracts ``{value, evidence, method}``:

* ``value``    — the event's human payload (a run/cell text, a rule condition, a comment anchor text).
* ``evidence`` — ``file`` + location + the raw structured attributes (hex colours, formulas, ids…).
* ``method``   — ``engine`` (xlsx/docx), ``kind`` (``conditional_format`` / ``cell`` / ``run`` /
  ``comment``), the coarse ``fill_color`` / ``font_color`` used by the composite predicate, and a
  ``source`` tag (``conditional`` / ``direct`` / ``highlight`` / ``shading``).
"""
from __future__ import annotations

import io
import os
import zipfile
from pathlib import Path
from typing import Any

from src.rag.corpus import FileRef, nfc
from src.rag.extract import office as _office
from src.rag.tools import contract
from src.rag.tools.contract import ContractError
from src.rag.tools.extract_tools import resolve_ref
from src.rag.tools.highlight_extract import (
    _COLOR_FAMILY, _WORD_HIGHLIGHT_NAMES, _norm_color_query, _office_bytes,
)
from src.rag.tools.profile import CorpusProfile

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

_ON = {"1", "true", "yes", "on"}

# Office formats whose OOXML semantic layer this tool reads deterministically.
_XLSX_EXTS = {"xlsx", "xlsm"}
_OFFICE_EXTS = _XLSX_EXTS | {"docx"}


def enabled() -> bool:
    """True when the FORMAT_EVENT extract face / tool is active (default OFF, opt-in like siblings)."""
    return os.getenv("RAG_FORMAT_EVENTS", "0").strip().lower() in _ON


# --------------------------------------------------------------------------- colour helpers
def _color_matches(item_color: str | None, target: str | None) -> bool:
    """True when no filter is set, or the item's coarse colour equals the target (青⊇水色 family-aware)."""
    if target is None:
        return True
    if item_color is None:
        return False
    ic = nfc(item_color).strip()
    if ic.lower() == target.strip().lower():
        return True
    return ic in _COLOR_FAMILY.get(target.strip(), frozenset())


def _event(value: Any, *, engine: str, kind: str, source: str,
           fill_color: str | None, font_color: str | None,
           **evidence: Any) -> dict[str, Any]:
    """Build one per-event contract ``{value, evidence, method}``."""
    return {
        "value": value,
        "evidence": {k: v for k, v in evidence.items() if v is not None},
        "method": {"engine": engine, "kind": kind, "source": source,
                   "fill_color": fill_color, "font_color": font_color},
    }


def fmt_notation(item: dict[str, Any]) -> str:
    """Render one event as a single machine-searchable ``__FMT__`` line.

    Only the deterministic, lexically-filterable predicate attributes are emitted (coarse colours, the
    background provenance, the style origin, and — for a conditional rule — its condition). Absent
    attributes are omitted so a filter like ``background=黄`` matches exactly.
    """
    m = item.get("method", {})
    ev = item.get("evidence", {})
    parts: list[str] = []
    if m.get("font_color"):
        parts.append(f"font.color={m['font_color']}")
    if m.get("fill_color"):
        parts.append(f"background={m['fill_color']}")
    src = m.get("source")
    if src and m.get("fill_color"):
        parts.append(f"background.source={src}")
    if ev.get("style_origin"):
        parts.append(f"style.origin={ev['style_origin']}")
    if m.get("kind") == "conditional_format" and ev.get("condition"):
        parts.append(f'rule="{ev["condition"]}"')
    if m.get("kind") == "comment":
        parts.append("comment=1")
    return "__FMT__ " + " ".join(parts) if parts else "__FMT__"


# =========================================================================== Excel
def _cf_condition_text(operator: str | None, formulas: Any) -> str:
    """Human-readable condition for a ``cellIs`` rule (e.g. ``セルの値 < -0.99``)."""
    ops = {"lessThan": "<", "lessThanOrEqual": "<=", "greaterThan": ">",
           "greaterThanOrEqual": ">=", "equal": "=", "notEqual": "!=",
           "between": "の範囲", "notBetween": "の範囲外", "containsText": "を含む"}
    vals = ", ".join(str(x) for x in (formulas or []))
    label = ops.get(operator or "", operator or "")
    if operator in ("between", "notBetween"):
        return f"セルの値が {vals} {label}".strip()
    return f"セルの値 {label} {vals}".strip()


def _cf_satisfying_cells(ws, sqref, operator: str | None,
                         formulas: Any) -> list[tuple[str, Any]]:
    """Numeric cells inside ``sqref`` that satisfy a ``cellIs`` rule, as ``(coordinate, value)``.

    openpyxl is not a calc engine (``data_only=True`` gives the *cached* value), so only self-contained
    numeric ``cellIs`` rules are evaluated; formula/expression rules are surfaced as the rule alone.
    """
    try:
        thr = [float(str(f)) for f in (formulas or [])]
    except (TypeError, ValueError):
        return []
    if not thr:
        return []

    def ok(x: float) -> bool:
        if operator == "lessThan":
            return x < thr[0]
        if operator == "lessThanOrEqual":
            return x <= thr[0]
        if operator == "greaterThan":
            return x > thr[0]
        if operator == "greaterThanOrEqual":
            return x >= thr[0]
        if operator == "equal":
            return x == thr[0]
        if operator == "notEqual":
            return x != thr[0]
        if operator == "between" and len(thr) >= 2:
            return thr[0] <= x <= thr[1]
        if operator == "notBetween" and len(thr) >= 2:
            return not (thr[0] <= x <= thr[1])
        return False

    out: list[tuple[str, Any]] = []
    for cr in sqref.ranges:
        min_col, min_row, max_col, max_row = cr.bounds
        for r in range(min_row, max_row + 1):
            for c in range(min_col, max_col + 1):
                v = ws.cell(r, c).value
                if isinstance(v, bool) or not isinstance(v, (int, float)):
                    continue
                if ok(float(v)):
                    out.append((ws.cell(r, c).coordinate, v))
    return out


def _dxf_fill_color(dxf) -> tuple[str | None, str | None]:
    """``(coarse_name, hex)`` of a differential-style fill (bgColor preferred, then fgColor)."""
    fill = getattr(dxf, "fill", None) if dxf else None
    if fill is None:
        return None, None
    for attr in ("bgColor", "fgColor"):
        c = getattr(fill, attr, None)
        if c is None:
            continue
        name = _office._excel_color_name(c)
        if name:
            return name, str(getattr(c, "rgb", "") or "") or None
    return None, None


def _dxf_font_color(dxf) -> tuple[str | None, str | None]:
    """``(coarse_name, hex)`` of a differential-style font colour."""
    font = getattr(dxf, "font", None) if dxf else None
    color = getattr(font, "color", None) if font is not None else None
    if color is None:
        return None, None
    return _office._excel_color_name(color), str(getattr(color, "rgb", "") or "") or None


def _dxf_border_sides(dxf) -> list[str]:
    """Which border sides a differential style sets (so a 'border rule' is surfaced, not dropped)."""
    border = getattr(dxf, "border", None) if dxf else None
    if border is None:
        return []
    sides = []
    for side in ("left", "right", "top", "bottom"):
        s = getattr(border, side, None)
        if s is not None and getattr(s, "style", None):
            sides.append(side)
    return sides


def _xlsx_cf_events(ws, file: str) -> list[dict[str, Any]]:
    """One FORMAT_EVENT per conditional-format rule on a worksheet (full dxf/priority/stop record)."""
    events: list[dict[str, Any]] = []
    cf = ws.conditional_formatting
    for rng in cf:
        for rule in cf[rng]:
            dxf = getattr(rule, "dxf", None)
            fill_name, fill_hex = _dxf_fill_color(dxf)
            font_name, font_hex = _dxf_font_color(dxf)
            border_sides = _dxf_border_sides(dxf)
            formulas = [str(f) for f in (getattr(rule, "formula", None) or [])]
            cond = _cf_condition_text(getattr(rule, "operator", None), formulas)
            sat = (_cf_satisfying_cells(ws, rng.sqref, rule.operator, rule.formula)
                   if getattr(rule, "type", None) == "cellIs" else [])
            events.append(_event(
                cond or f"{rule.type} rule", engine="xlsx", kind="conditional_format",
                source="conditional", fill_color=fill_name, font_color=font_name,
                file=file, sheet=ws.title, range=str(rng.sqref),
                rule_type=getattr(rule, "type", None), operator=getattr(rule, "operator", None),
                formula=formulas or None, priority=getattr(rule, "priority", None),
                stop_if_true=getattr(rule, "stopIfTrue", None), dxf_id=getattr(rule, "dxfId", None),
                dxf_fill_hex=fill_hex, dxf_font_hex=font_hex,
                dxf_border=border_sides or None, condition=cond or None,
                n_matched=len(sat) or None,
                matched_cells=[c for c, _ in sat][:50] or None,
                matched_values=[v for _, v in sat][:50] or None))
    return events


def _xlsx_direct_events(ws, file: str) -> list[dict[str, Any]]:
    """Direct-style cells carrying a fill and/or font colour (Level 1) — for composite predicates."""
    events: list[dict[str, Any]] = []
    for row in ws.iter_rows():
        for cell in row:
            if cell.value is None:
                continue
            fill = cell.fill
            fill_name = None
            if fill is not None and getattr(fill, "patternType", None):
                fill_name = _office._excel_color_name(getattr(fill, "fgColor", None))
            font = cell.font
            font_name = _office._excel_color_name(getattr(font, "color", None)) if font else None
            if not (fill_name or font_name):
                continue
            events.append(_event(
                str(cell.value), engine="xlsx", kind="cell", source="direct",
                fill_color=fill_name, font_color=font_name,
                file=file, sheet=ws.title, cell=cell.coordinate,
                row=cell.row, column=cell.column, style_origin="direct"))
    return events


def _xlsx_events(ref: FileRef, data: bytes | None) -> list[dict[str, Any]]:
    """All Excel FORMAT_EVENTs: conditional-format rules (Level 2) + direct-style cells (Level 1)."""
    import openpyxl

    src = io.BytesIO(data) if data else str(ref.path)
    wb = openpyxl.load_workbook(src, data_only=True)
    file = ref.rel or str(ref.path)
    events: list[dict[str, Any]] = []
    for ws in wb.worksheets:
        events.extend(_xlsx_cf_events(ws, file))
        events.extend(_xlsx_direct_events(ws, file))
    return events


# =========================================================================== Word
def _resolve_font_color(color_el) -> tuple[str | None, str | None, str | None]:
    """``(coarse_name, resolved_hex, theme_token)`` from a ``w:color`` element.

    A direct ``w:val`` hex resolves through the same classifier as fills; a ``w:themeColor`` (whose RGB
    depends on the document theme) is reported by *token* and fails closed for the coarse name — never
    guessed — so a composite ``font.color==赤`` predicate stays deterministic.
    """
    if color_el is None:
        return None, None, None
    val = color_el.get(_W + "val")
    theme = color_el.get(_W + "themeColor")
    if val and val.lower() not in ("auto",):
        return _office._color_name(val), val.upper(), theme
    return None, None, theme


def _shd_fill(shd_el) -> tuple[str | None, str | None]:
    """``(coarse_name, hex)`` of a ``w:shd`` shading element's fill (drops auto / white)."""
    if shd_el is None:
        return None, None
    fill = shd_el.get(_W + "fill")
    if fill and fill.lower() not in ("auto", "ffffff"):
        return _office._color_name(fill), fill.upper()
    return None, None


def _run_highlight_name(rpr) -> str | None:
    """Coarse colour of a run's ``w:highlight`` (the ST_HighlightColor keyword palette)."""
    if rpr is None:
        return None
    hl = rpr.find(_W + "highlight")
    if hl is None:
        return None
    val = (hl.get(_W + "val") or "").lower()
    if val and val not in ("none", "auto"):
        return _WORD_HIGHLIGHT_NAMES.get(val, val)
    return None


def _docx_style_events(ref: FileRef, data: bytes | None) -> list[dict[str, Any]]:
    """Per-run effective-style events: separate direct/theme font colour, highlight, and shadings.

    The **effective background** follows the OOXML display rule that a text *highlight* renders over run
    ``w:shd`` shading (Microsoft OOXML docs), which over paragraph/cell shading — so a composite
    ``background`` predicate reflects what a reader sees. Paragraph and table-cell shading are threaded
    down as context so a run inheriting its background from the paragraph/cell still exposes it.
    """
    import docx

    d = docx.Document(io.BytesIO(data) if data else str(ref.path))
    file = ref.rel or str(ref.path)
    events: list[dict[str, Any]] = []

    def emit_paragraph(para, *, cell_shd: tuple[str | None, str | None], location: str) -> None:
        ppr = para._p.find(_W + "pPr")
        para_shd = _shd_fill(ppr.find(_W + "shd")) if ppr is not None else (None, None)
        for run in para.runs:
            rpr = run._r.find(_W + "rPr")
            text = run.text.strip()
            if not text:
                continue
            font_name, font_hex, theme = _resolve_font_color(
                rpr.find(_W + "color") if rpr is not None else None)
            highlight = _run_highlight_name(rpr)
            run_shd_name, run_shd_hex = _shd_fill(rpr.find(_W + "shd") if rpr is not None else None)
            # effective background: highlight > run shading > paragraph shading > cell shading
            bg_name, bg_source = None, "none"
            if highlight:
                bg_name, bg_source = highlight, "highlight"
            elif run_shd_name:
                bg_name, bg_source = run_shd_name, "shading"
            elif para_shd[0]:
                bg_name, bg_source = para_shd[0], "paragraph_shading"
            elif cell_shd[0]:
                bg_name, bg_source = cell_shd[0], "cell_shading"
            if not (font_name or bg_name):
                continue
            # style origin: a direct rPr colour/highlight is "direct"; a colour inherited from the
            # paragraph/cell shading with no direct run colour is "inherited"; otherwise "style".
            direct = bool((rpr is not None and rpr.find(_W + "color") is not None)
                          or highlight or run_shd_name)
            style_origin = "direct" if direct else "inherited"
            events.append(_event(
                text, engine="docx", kind="run",
                source=bg_source if bg_name else "direct",
                fill_color=bg_name, font_color=font_name,
                file=file, paragraph=location,
                direct_font_hex=font_hex, theme_font_color=theme, resolved_font_hex=font_hex,
                highlight=highlight, run_shading=run_shd_hex,
                paragraph_shading=para_shd[1], cell_shading=cell_shd[1],
                effective_background=bg_name, background_source=bg_source if bg_name else None,
                style_origin=style_origin))

    for pi, para in enumerate(d.paragraphs, 1):
        emit_paragraph(para, cell_shd=(None, None), location=f"para:{pi}")
    for ti, table in enumerate(d.tables, 1):
        for ri, trow in enumerate(table.rows, 1):
            for ci, tcell in enumerate(trow.cells, 1):
                tcpr = tcell._tc.find(_W + "tcPr")
                cell_shd = _shd_fill(tcpr.find(_W + "shd")) if tcpr is not None else (None, None)
                for pi, para in enumerate(tcell.paragraphs, 1):
                    emit_paragraph(para, cell_shd=cell_shd,
                                   location=f"表{ti}:R{ri}C{ci}:para:{pi}")
    return events


# --------------------------------------------------------------------------- docx comments (XML-direct)
def _docx_comment_events(ref: FileRef, data: bytes | None) -> list[dict[str, Any]]:
    """Every docx comment with its **anchor text**, via a deterministic ``document.xml`` ⇄ ``comments.xml``
    correspondence (python-docx 1.1.x exposes no comment API, so the XML-direct path is primary).

    The anchor text is the run text enclosed by the matching ``w:commentRangeStart`` /
    ``w:commentRangeEnd`` (by ``w:id``); the comment body/author/date come from ``comments.xml``; the
    anchor paragraph id is the ``w14:paraId`` (or 1-based index) of the paragraph holding the range start.
    """
    from lxml import etree

    src = io.BytesIO(data) if data else str(ref.path)
    file = ref.rel or str(ref.path)
    try:
        zf = zipfile.ZipFile(src)
    except Exception:
        return []
    with zf:
        names = set(zf.namelist())
        if "word/comments.xml" not in names or "word/document.xml" not in names:
            return []
        try:
            comments_xml = etree.fromstring(zf.read("word/comments.xml"))
            document_xml = etree.fromstring(zf.read("word/document.xml"))
        except Exception:
            return []

    W14 = "{http://schemas.microsoft.com/office/word/2010/wordml}"
    # comment body / author / date, keyed by w:id
    meta: dict[str, dict[str, Any]] = {}
    for c in comments_xml.findall(_W + "comment"):
        cid = c.get(_W + "id")
        if cid is None:
            continue
        body = "".join(t.text or "" for t in c.iter(_W + "t")).strip()
        meta[cid] = {"author": c.get(_W + "author"), "date": c.get(_W + "date"), "text": body}

    # anchor text + anchor paragraph id per comment id, in document order
    events: list[dict[str, Any]] = []
    active: dict[str, list[str]] = {}
    anchor_para: dict[str, str] = {}
    para_index = 0
    for el in document_xml.iter():
        tag = el.tag
        if tag == _W + "p":
            para_index += 1
            cur_para = el.get(W14 + "paraId") or f"index:{para_index}"
        elif tag == _W + "commentRangeStart":
            cid = el.get(_W + "id")
            if cid is not None:
                active[cid] = []
                anchor_para[cid] = locals().get("cur_para", f"index:{para_index}")
        elif tag == _W + "commentRangeEnd":
            cid = el.get(_W + "id")
            if cid is not None and cid in active:
                anchor = "".join(active.pop(cid)).strip()
                info = meta.get(cid, {})
                events.append(_event(
                    anchor or info.get("text", ""), engine="docx", kind="comment", source="comment",
                    fill_color=None, font_color=None,
                    file=file, comment_id=cid, author=info.get("author"),
                    timestamp=info.get("date"), anchor_text=anchor or None,
                    anchor_paragraph_id=anchor_para.get(cid), comment_text=info.get("text") or None))
        elif tag == _W + "t" and el.text:
            for cid in active:
                active[cid].append(el.text)
    # comments whose range markers are absent (some producers only emit a commentReference): still surface
    # the comment body so idx49-style "コメントを抽出" is answerable even without a resolvable anchor.
    emitted = {ev["evidence"].get("comment_id") for ev in events}
    for cid, info in meta.items():
        if cid in emitted or not info.get("text"):
            continue
        events.append(_event(
            info["text"], engine="docx", kind="comment", source="comment",
            fill_color=None, font_color=None,
            file=file, comment_id=cid, author=info.get("author"), timestamp=info.get("date"),
            anchor_text=None, anchor_paragraph_id=None, comment_text=info.get("text")))
    return events


def _docx_events(ref: FileRef, data: bytes | None) -> list[dict[str, Any]]:
    """All docx FORMAT_EVENTs: effective-style runs + comment anchors."""
    return _docx_style_events(ref, data) + _docx_comment_events(ref, data)


# =========================================================================== dispatch / public API
def _all_events(ref: FileRef, profile: CorpusProfile | None) -> list[dict[str, Any]]:
    """Every deterministic FORMAT_EVENT for an Office file (no vision, no rendering)."""
    data = _office_bytes(ref, profile)
    if ref.ext in _XLSX_EXTS:
        return _xlsx_events(ref, data)
    if ref.ext == "docx":
        return _docx_events(ref, data)
    raise ContractError(
        f"unsupported file for format events: .{ref.ext} (expected xlsx/xlsm/docx)")


def format_events(file: str | Path | FileRef, *, fill: str | None = None,
                  font_color: str | None = None, source: str | None = None,
                  kind: str | None = None,
                  profile: CorpusProfile | None = None) -> dict[str, Any]:
    """Enumerate deterministic OOXML FORMAT_EVENTs in an Excel/Word file, filterable by a composite predicate.

    Parameters
    ----------
    file
        A :class:`FileRef`, absolute path, or NFC corpus-relative name/rel (resolved NFC-aware).
    fill
        Optional **background** colour filter (coarse Japanese name / synonym; 青⊇水色 family-aware).
        Matches a docx run's effective background or an Excel fill / conditional-format dxf fill.
    font_color
        Optional **font** colour filter (same vocabulary). Combined with ``fill`` as a logical AND — the
        idx16 「黄色ハイライトかつ赤字」 composite predicate.
    source
        Optional filter on the event's provenance tag (``conditional`` / ``direct`` / ``highlight`` /
        ``shading`` / ``comment``).
    kind
        Optional filter on the event kind (``conditional_format`` / ``cell`` / ``run`` / ``comment``).
    profile
        Optional :class:`CorpusProfile` for transparent decryption of an encrypted Office file.

    ``value`` is a list of per-event contracts ``{value, evidence, method}`` in document order.
    Raises :class:`ContractError` for an unsupported extension or an unreadable file.
    """
    ref = resolve_ref(file)
    fill_t = _norm_color_query(fill)
    font_t = _norm_color_query(font_color)

    from src.rag.index import structure_store
    events = structure_store.lookup_format_events(ref)
    if events is None:
        events = _all_events(ref, profile)

    def keep(ev: dict[str, Any]) -> bool:
        m = ev["method"]
        if fill_t is not None and not _color_matches(m.get("fill_color"), fill_t):
            return False
        if font_t is not None and not _color_matches(m.get("font_color"), font_t):
            return False
        if source is not None and m.get("source") != source:
            return False
        if kind is not None and m.get("kind") != kind:
            return False
        return True

    items = [ev for ev in events if keep(ev)]
    kinds = sorted({ev["method"]["kind"] for ev in items})
    return contract.make(
        items,
        engine="format_events",
        evidence={"file": ref.rel or str(ref.path), "ext": ref.ext,
                  "n_events": len(items), "kinds": kinds},
        fill_filter=fill_t, font_color_filter=font_t,
    )


def comment_anchors(file: str | Path | FileRef, *,
                    profile: CorpusProfile | None = None) -> list[str]:
    """Convenience: the ordered, de-duplicated docx comment anchor texts (bare list, not a contract).

    For 「コメントがついている部分をそのまま抽出」 (idx49): each comment's anchor text (the span the
    comment is attached to), falling back to the comment body when no anchor range is resolvable.
    Returns ``[]`` when the file carries no comments.
    """
    result = format_events(file, kind="comment", profile=profile)
    seen: set[str] = set()
    out: list[str] = []
    for item in result["value"]:
        text = str(item["value"]).strip()
        if text and text not in seen:
            seen.add(text)
            out.append(text)
    return out
